from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt


BASE_DIR = Path("/Users/janpetervisser/Development/third/docs")
PRODUCT_NAME = "Inspannings Monitor"
DATE_TEXT = "18 april 2026"
POSITIONING = "Wellness/self-management"
HOSTING = "Vercel"
DATABASE = "Supabase PostgreSQL"
AUTH = "Supabase Auth"


def init_doc(title_text: str, subtitle_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos"
    styles["Title"].font.size = Pt(22)
    styles["Subtitle"].font.size = Pt(11)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(title_text)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(subtitle_text)
    doc.add_paragraph("")
    return doc


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)
    run.append(run_props)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def p(doc: Document, text: str = "", style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def bullets(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, headers, rows) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph("")


def set_footer(doc: Document, text: str) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = text


def build_testplan() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Testplan v0.2",
        f"Risicogebaseerde teststrategie, traceability, cybersecurity en QMS-fundament\n{DATE_TEXT}",
    )

    # ── 1. Documentdoel ──────────────────────────────────────────────────────
    p(doc, "1. Documentdoel", "Heading 1")
    p(
        doc,
        f"Dit document beschrijft hoe {PRODUCT_NAME} getest wordt. De strategie is verschoven van "
        "'werkt de feature?' naar 'is het risico beheerst?'. "
        "Dat onderscheid is relevant zelfs voor een wellness-first product: de app verwerkt gezondheidsgerelateerde "
        "gegevens van mensen met chronische vermoeidheid, en een fout in de budgetberekening of een lek in de "
        "datatoegang kan directe gevolgen hebben voor het vertrouwen en zelfmanagement van de gebruiker. "
        "Dit document combineert risicoanalyse (ISO 14971-principe), verificatie versus validatie, "
        "een traceability matrix, cybersecurity testing (NEN 7510), ISO 13485-voorbereiding "
        "en de praktische testpiramide met tooling.",
    )

    # ── 2. Van feature-test naar risicobeheer ────────────────────────────────
    p(doc, "2. Teststrategie: van feature-test naar risicobeheer", "Heading 1")
    p(
        doc,
        "Traditionele teststrategie stelt de vraag: 'doet de code wat de developer verwacht?' "
        "Een risicogebaseerde aanpak stelt de vraag: 'wat zijn de gevolgen als deze functie faalt, en zijn die gevolgen acceptabel?' "
        "De testinspanning wordt verdeeld op basis van de risicoscore van een functie, niet op basis van complexiteit of backlog-prioriteit.",
    )
    table(
        doc,
        ["Risicoscore", "Betekenis", "Vereiste testdekking"],
        [
            ["Kritiek", "Fout leidt tot verkeerde zelfmanagementbeslissing of datadiefstal", "100% — unit, integratie én E2E verplicht"],
            ["Hoog", "Fout leidt tot onjuiste weergave of verlies van gebruikersdata", "Volledige unit- en integratietestdekking, E2E aanbevolen"],
            ["Middel", "Fout leidt tot verminderd gebruiksgemak of niet-kritieke weergavefout", "Minimaal unit tests op grenzen, handmatige verificatie"],
            ["Laag", "Cosmetische of zeldzame randgevallen zonder impact op data of beslissingen", "Handmatige QA voldoende"],
        ],
    )

    # ── 3. Verificatie versus validatie ──────────────────────────────────────
    p(doc, "3. Verificatie en validatie", "Heading 1")
    p(
        doc,
        "Dit onderscheid is cruciaal bij gezondheidsgerelateerde software en vormt de basis voor "
        "MDR-documentatieplicht en ISO 13485-conformiteit zodra het product richting een medisch track gaat. "
        "Ook voor de huidige wellness-first versie geldt dit als kwalitatief fundament.",
    )
    table(
        doc,
        ["Type", "Centrale vraag", "Activiteiten", "Wie"],
        [
            [
                "Verificatie",
                "Hebben we het product goed gebouwd?",
                "Unit tests, integratietests, RLS-tests, code review, statische analyse, traceability matrix",
                "Engineers",
            ],
            [
                "Validatie",
                "Hebben we het juiste product gebouwd voor de gebruiker?",
                "Usability tests met echte gebruikers, acceptatietests op functionele requirements, klinische evaluatie (toekomstig medisch track)",
                "Product, UX, (toekomstig) klinisch evaluator",
            ],
        ],
    )
    p(
        doc,
        "Validatie is niet hetzelfde als E2E-testen. Een geautomatiseerde Playwright-test verifieert dat de app "
        "technisch correct werkt, maar valideert niet of de gebruiker de energieslider begrijpt of de juiste "
        "beslissing neemt op basis van weergegeven informatie. Usability tests met echte gebruikers zijn nodig "
        "voor validatie — dit is een geplande activiteit vóór launch (ST-801).",
    )

    # ── 4. ISO 13485 en QMS-voorbereiding ────────────────────────────────────
    p(doc, "4. ISO 13485 en kwaliteitsmanagementsysteem (QMS)", "Heading 1")
    p(
        doc,
        "ISO 13485 is de internationale norm voor kwaliteitsmanagementsystemen voor medische hulpmiddelen. "
        f"In de huidige wellness-first fase is {PRODUCT_NAME} geen medisch hulpmiddel en is ISO 13485-certificering "
        "niet vereist. De norm wordt hier als leidraad gebruikt om de kwaliteitsborging al in de goede richting "
        "op te zetten, zodat de drempel naar een eventuele medische track zo laag mogelijk blijft.",
    )

    p(doc, "4.1 Relevante ISO 13485-vereisten als leidraad", "Heading 2")
    table(
        doc,
        ["ISO 13485-element", "Toepassing in wellness-first fase", "Status"],
        [
            ["Documentbeheer (par. 4.2)", "Alle specificaties, testplannen en besluitlogs worden beheerd en voorzien van versienummer en datum", "Ingericht via docs/ en git"],
            ["Risicomanagement (par. 7.1, ref. ISO 14971)", "Risicoanalyse per functie uitgevoerd (zie sectie 5)", "In dit document"],
            ["Software-ontwikkelingsproces (par. 7.3)", "Backlog, epics, definition of done en code review zijn aanwezig", "Ingericht via Linear en GitHub"],
            ["Verificatie en validatie (par. 7.3.6 / 7.3.7)", "Testpiramide met traceability matrix per requirement", "In dit document"],
            ["Traceability (par. 7.3.4)", "Traceability matrix koppelt elke FR aan een test en resultaat", "Sectie 6 van dit document"],
            ["Correctieve maatregelen / CAPA (par. 8.5)", "Bugs worden bijgehouden in Linear; ernstige fouten krijgen een root-cause analyse", "Aanbevolen werkwijze"],
            ["Interne audit (par. 8.2.2)", "Periodieke review van testresultaten en traceability matrix vóór elke release", "Aanbevolen als releasepoort"],
        ],
    )

    p(doc, "4.2 Wanneer is formele ISO 13485-certificering nodig?", "Heading 2")
    p(
        doc,
        "Formele certificering is verplicht wanneer het product wordt geclassificeerd als medisch hulpmiddel "
        "onder de EU Medical Device Regulation (MDR 2017/745). "
        "De decision gate hiervoor staat beschreven in het Roadmap-document (doc 04, Gate D en E). "
        "Zolang het product binnen de wellness-positionering blijft, is certificering niet vereist maar is "
        "de gestructureerde aanpak uit dit testplan voldoende als kwaliteitsborging.",
    )
    bullets(
        doc,
        [
            "Stel nu al een documentregister in (versies, goedkeuringen, archief) — dit is de kern van elk QMS.",
            "Documenteer afwijkingen en bugs als 'non-conformities' in Linear met root-cause en correctieve actie.",
            "Voer vóór elke release een interne review uit van de traceability matrix en de testresultaten.",
            "Zodra het product richting medisch track gaat: stel een formeel QMS-handboek op en start een gap-analyse tegen ISO 13485.",
        ],
    )

    # ── 5. Risicoanalyse per functie (ISO 14971) ─────────────────────────────
    p(doc, "5. Risicoanalyse per functie (ISO 14971-principe)", "Heading 1")
    p(
        doc,
        "Voor elke kritieke functie is bepaald wat de ernstigste consequentie van een fout is, "
        "wat de risicoscore is en welke testdekking verplicht is. "
        "Risico = kans op optreden × ernst van de gevolgen. "
        "Kans wordt bepaald door code-complexiteit en het ontbreken van type-veiligheid of validatie.",
    )
    table(
        doc,
        ["Functie / module", "Ernstigste fout", "Ernst", "Kans zonder tests", "Risicoscore", "Vereiste dekking"],
        [
            [
                "Budgetberekening (score → energyLevel + dailyBudget)",
                "Gebruiker krijgt verkeerd budget; plant meer dan verantwoord is",
                "Hoog",
                "Middel",
                "Kritiek",
                "100% unit tests op alle grenswaarden + integratietest op opgeslagen output",
            ],
            [
                "RLS-beleid (owner-only datatoegang)",
                "Gebruiker A leest of overschrijft data van gebruiker B",
                "Kritiek",
                "Middel",
                "Kritiek",
                "100% pgTAP-tests op alle tabellen, alle operaties, alle rollen",
            ],
            [
                "Authenticatie en sessiebeheer",
                "Niet-ingelogde gebruiker krijgt toegang tot dashboard of data",
                "Hoog",
                "Laag",
                "Hoog",
                "E2E-test op elke beveiligde route; integratietest op getAuthState()",
            ],
            [
                "Zod-invoervalidatie (server actions)",
                "Ongeldige of kwaadaardige invoer bereikt de database",
                "Hoog",
                "Middel",
                "Hoog",
                "Unit tests op schema-grenzen; integratietest op server action met ongeldige invoer",
            ],
            [
                "Insightregels en weekpatronen",
                "Gebruiker trekt verkeerde conclusie op basis van onjuist patroon",
                "Middel",
                "Middel",
                "Hoog",
                "Unit tests op aggregatiefuncties; datadrempellogica getest op minimum-threshold",
            ],
            [
                "Reflectieprompt-planning (T+1/T+2 job)",
                "Dubbele of gemiste prompts",
                "Laag",
                "Middel",
                "Middel",
                "Integratietest op idempotentie van joblogica",
            ],
            [
                "Navigatie-sanitisatie (open redirect)",
                "Gebruiker wordt doorgestuurd naar externe kwaadaardige URL",
                "Hoog",
                "Laag",
                "Hoog",
                "Unit tests inclusief double-slash en externe URL-patronen",
            ],
            [
                "Copy en insighttekst (wellness vs. medisch)",
                "Regulatoire interpretatie als medisch hulpmiddel",
                "Hoog (regulatoir)",
                "Middel",
                "Hoog",
                "Handmatige copyreview vóór elke release; geautomatiseerde woordlijstcheck overwegen",
            ],
        ],
    )

    # ── 6. Traceability matrix ───────────────────────────────────────────────
    p(doc, "6. Traceability matrix", "Heading 1")
    p(
        doc,
        "De traceability matrix koppelt elke functionele requirement (FR-ID) aan de tests die aantonen dat de eis "
        "is geverifieerd. Dit is een kernvereiste voor MDR-conformiteit, ISO 13485 (par. 7.3.4) en voor "
        "auditeerbare kwaliteitsborging. De matrix wordt bij elke release bijgewerkt met het testresultaat. "
        "Mislukte tests blokkeren de release. Overgeslagen tests vereisen een expliciete risicoafweging.",
    )
    table(
        doc,
        ["Requirement ID", "Omschrijving (kort)", "Testtype", "Test ID / bestand", "Risicoscore", "Resultaat"],
        [
            ["FR-CHK-001", "Check-in opslaan met energiescore en slaapkwaliteit", "E2E", "e2e/checkin.spec.ts — happy path", "Middel", "—"],
            ["FR-CHK-002", "Dagbudget afleiden uit energiescore", "Unit + Integratie", "lib/checkin/__tests__/budget.test.ts", "Kritiek", "—"],
            ["FR-PLAN-001", "Activiteit plannen met verplichte velden", "E2E", "e2e/planning.spec.ts — aanmaken", "Middel", "—"],
            ["FR-PLAN-002", "Lopend totaal bijwerken na mutatie", "Integratie + E2E", "lib/planning/__tests__/meter.test.ts", "Hoog", "—"],
            ["FR-PLAN-003", "Niet-blokkerende waarschuwing bij overschrijding", "E2E", "e2e/planning.spec.ts — budget overschrijden", "Middel", "—"],
            ["FR-ACT-001", "Activiteit als uitgevoerd markeren", "E2E", "e2e/planning.spec.ts — uitgevoerd", "Middel", "—"],
            ["FR-ACT-002", "Activiteit als geskipt markeren met reden", "E2E", "e2e/planning.spec.ts — geskipt", "Middel", "—"],
            ["FR-ACT-003", "Activiteit als aangepast markeren", "E2E", "e2e/planning.spec.ts — aangepast", "Middel", "—"],
            ["FR-ACT-005", "Ongeplande activiteit toevoegen", "Integratie + E2E", "lib/planning/__tests__/service.test.ts", "Middel", "—"],
            ["FR-DAY-001", "Gepland versus uitgevoerd tonen in dagoverzicht", "E2E", "e2e/planning.spec.ts — dagoverzicht", "Middel", "—"],
            ["FR-WEEK-001", "Weekoverzicht met gemiddelde energie en adherence", "Unit + Integratie", "lib/insights/__tests__/week.test.ts", "Hoog", "—"],
            ["FR-INS-001", "Inzicht alleen tonen bij minimale data", "Unit", "lib/insights/__tests__/thresholds.test.ts", "Hoog", "—"],
            ["FR-REM-001", "Reflectieprompts per gebruiker aan/uit", "Integratie", "lib/reflection/__tests__/service.test.ts", "Middel", "—"],
            ["FR-SET-001", "Instellingen opslaan en direct actief", "E2E", "e2e/settings.spec.ts", "Middel", "—"],
            ["SEC-001", "TLS op alle communicatie", "Handmatig / infra", "SSL Labs check op productiedomain", "Hoog", "—"],
            ["SEC-004", "Rate limiting op auth-routes", "Integratie", "lib/auth/__tests__/ratelimit.test.ts", "Hoog", "—"],
            ["SEC-RLS-001", "Owner-only SELECT op profiles", "pgTAP", "supabase/tests/test_profiles_rls.sql", "Kritiek", "—"],
            ["SEC-RLS-002", "Owner-only SELECT op user_settings", "pgTAP", "supabase/tests/test_user_settings_rls.sql", "Kritiek", "—"],
            ["SEC-RLS-003", "Owner-only SELECT op morning_check_ins", "pgTAP", "supabase/tests/test_morning_check_ins_rls.sql", "Kritiek", "—"],
            ["SEC-RLS-004", "Owner-only SELECT op activities", "pgTAP", "supabase/tests/test_activities_rls.sql", "Kritiek", "—"],
            ["SAFE-001", "Geen medische taal in UI-copy", "Handmatig copyreview", "Checklist ST-803", "Hoog", "—"],
            ["SAFE-003", "Inzichten tonen minimale-dataguardrail", "Unit", "lib/insights/__tests__/thresholds.test.ts", "Hoog", "—"],
        ],
    )

    # ── 7. Cybersecurity testing (NEN 7510) ──────────────────────────────────
    p(doc, "7. Cybersecurity testing (NEN 7510)", "Heading 1")
    p(
        doc,
        "NEN 7510 is de Nederlandse norm voor informatiebeveiliging in de zorg. "
        f"Hoewel {PRODUCT_NAME} een wellness-product is, verwerkt de app gezondheidsgerelateerde persoonsgegevens. "
        "De NEN 7510-baseline wordt als toetssteen gebruikt om privacyrisico's te beheersen en de "
        "drempel naar een toekomstige medische track te verlagen.",
    )

    p(doc, "7.1 Encryptie en datatransport", "Heading 2")
    table(
        doc,
        ["Eis", "Norm", "Testmethode", "Acceptatiecriterium"],
        [
            ["TLS 1.2 of hoger op alle routes", "NEN 7510 / SEC-001", "SSL Labs op productiedomain", "A-rating, geen TLS 1.0/1.1"],
            ["Data at rest versleuteld (AES-256)", "NEN 7510 / SEC-002", "Supabase-dashboard encryptie-instellingen", "AES-256 bevestigd"],
            ["Geen gevoelige data in URL-parameters", "OWASP / NEN 7510", "Handmatige review alle redirects", "Geen tokens of gezondheidsdata in URL"],
            ["HTTPS-only, geen mixed content", "SEC-001", "Content-Security-Policy header check", "Geen HTTP-requests in productie"],
            ["Veilige cookie-attributen", "OWASP Session Management", "Browser-devtools of Playwright-test", "Secure, HttpOnly, SameSite aanwezig"],
        ],
    )

    p(doc, "7.2 Authenticatie en toegangscontrole", "Heading 2")
    table(
        doc,
        ["Eis", "Norm", "Testmethode", "Acceptatiecriterium"],
        [
            ["Brute-force bescherming op login", "NEN 7510 / SEC-004", "Integratietest: >10 snelle pogingen triggert rate limit", "429-respons of vertraging"],
            ["Sessie vervalt na inactiviteit", "NEN 7510 / SEC-003", "Handmatig: sessie na 24 uur controleren", "Gebruiker moet opnieuw inloggen"],
            ["Geen sessietokens in localStorage", "OWASP", "Browser-devtools na login", "Geen tokens in localStorage"],
            ["Beveiligde routes zonder sessie", "SEC-003", "Playwright-test: dashboard zonder cookie", "Redirect naar /login"],
            ["Owner-only datatoegang (RLS)", "NEN 7510 / SEC-002", "pgTAP-tests op alle tabellen", "Zie traceability SEC-RLS-001 t/m 004"],
        ],
    )

    p(doc, "7.3 Penetratietest", "Heading 2")
    table(
        doc,
        ["Testvorm", "Scope", "Timing", "Uitvoerder"],
        [
            ["Geautomatiseerde OWASP Top 10 scan", "Alle publieke en beveiligde routes", "Vóór launch R1", "OWASP ZAP of Burp Suite Community"],
            ["Handmatig: SQL-injectie via formulieren", "Alle invoervelden die server actions aanroepen", "Vóór launch R1", "Engineer of security reviewer"],
            ["Handmatig: IDOR (cross-user data access)", "Alle calls met user-specifieke IDs", "Vóór launch R1", "Engineer"],
            ["Formele pentest door externe partij", "Volledige applicatie", "Vóór medische track", "Gecertificeerde pentest-partij"],
        ],
    )
    p(
        doc,
        "IDOR-testprocedure: log in als gebruiker A, kopieer een record-ID (bijv. activity ID), "
        "log in als gebruiker B en probeer dat record op te halen of te muteren via directe API-aanroep. "
        "Verwacht resultaat: 404 of 403, nooit de data van gebruiker A.",
    )

    p(doc, "7.4 Security headers", "Heading 2")
    table(
        doc,
        ["Header", "Aanbevolen waarde", "Testmethode"],
        [
            ["Content-Security-Policy", "Strikte policy, inline scripts beperkt", "securityheaders.com"],
            ["Strict-Transport-Security", "max-age=31536000; includeSubDomains", "curl -I op productiedomain"],
            ["X-Frame-Options", "DENY of SAMEORIGIN", "securityheaders.com"],
            ["X-Content-Type-Options", "nosniff", "securityheaders.com"],
            ["Referrer-Policy", "strict-origin-when-cross-origin", "securityheaders.com"],
        ],
    )
    p(doc, "Security headers worden geconfigureerd in next.config.ts via de headers()-functie.", "Normal")

    p(doc, "7.5 Logging en auditability (NEN 7510)", "Heading 2")
    table(
        doc,
        ["Te loggen event", "Minimale informatie", "Testmethode"],
        [
            ["Mislukte loginpoging", "Tijdstip, IP, e-mailadres (geanonimiseerd)", "Integratietest: mislukte login triggert logentry"],
            ["Succesvolle login", "Tijdstip, userId", "Integratietest: succesvolle login triggert logentry"],
            ["Accountverwijdering", "Tijdstip, userId", "Handmatige verificatie"],
            ["Sessie-timeout", "Tijdstip, userId", "Handmatige verificatie"],
        ],
    )

    # ── 8. Testpiramide en tooling ───────────────────────────────────────────
    p(doc, "8. Testpiramide en tooling", "Heading 1")
    table(
        doc,
        ["Laag", "Wat wordt getest", "Framework", "Wanneer"],
        [
            ["Unit", "Pure functies, berekeningslogica, Zod-schema's, hulpfuncties", "Vitest", "Bij elke commit"],
            ["Integratie", "Servicelaag, server actions, Supabase-queries", "Vitest + echte Supabase", "Bij elke commit"],
            ["Database / RLS", "RLS-beleid direct in PostgreSQL", "pgTAP via supabase test db", "Bij elke commit"],
            ["End-to-end", "Volledige gebruikersflows in echte browser", "Playwright", "Bij PR naar main"],
            ["Cybersecurity", "OWASP Top 10, headers, encryptie, IDOR", "ZAP + handmatig", "Vóór elke release"],
            ["Handmatig / validatie", "Usability, toegankelijkheid, copy, regressie", "Checklist", "Vóór elke release"],
        ],
    )

    p(doc, "8.1 Vitest — unit en integratietests", "Heading 2")
    p(doc, "npm install -D vitest @vitejs/plugin-react jsdom @testing-library/react @testing-library/dom vite-tsconfig-paths", "Normal")
    p(doc, "npx vitest                                    — watch mode", "Normal")
    p(doc, "npx vitest run                                — eenmalig alle tests", "Normal")
    p(doc, "npx vitest run lib/checkin/__tests__/budget.test.ts  — één bestand", "Normal")

    p(doc, "8.2 Playwright — end-to-end tests", "Heading 2")
    p(
        doc,
        "Authenticatie verloopt programmatisch via de Supabase Auth REST API (global-setup.ts). "
        "Het token wordt eenmalig opgehaald en hergebruikt als cookie-state voor alle tests.",
    )
    p(doc, "npm install -D @playwright/test && npx playwright install", "Normal")
    p(doc, "npx playwright test                           — alle E2E-tests", "Normal")
    p(doc, "npx playwright test --workers=1               — bij Supabase connection-limiet in CI", "Normal")

    p(doc, "8.3 Zod — runtime-validatie en testschema's", "Heading 2")
    p(doc, "npm install zod && npm install -D zod-fixture", "Normal")
    p(
        doc,
        "Zod-schema's in lib/*/schemas.ts worden hergebruikt in server actions én client-componenten. "
        "zod-fixture genereert automatisch testfixtures vanuit het schema.",
    )

    p(doc, "8.4 pgTAP — RLS en database-tests", "Heading 2")
    p(doc, "supabase test db     — voert alle .sql-testbestanden in supabase/tests/ uit", "Normal")
    p(
        doc,
        "pgTAP draait direct in PostgreSQL op hetzelfde uitvoerniveau als productieverkeer. "
        "Dit is de enige methode die RLS-gedrag betrouwbaar verifieert.",
    )

    # ── 9. Unit tests ────────────────────────────────────────────────────────
    p(doc, "9. Unit tests", "Heading 1")

    p(doc, "9.1 Budgetberekening (risicoscore: Kritiek)", "Heading 2")
    table(
        doc,
        ["Testgeval", "Input", "Verwacht resultaat"],
        [
            ["Minimale score", "energyScore = 1", "energyLevel = 'very_low', dailyBudget = minimumwaarde"],
            ["Maximale score", "energyScore = 10", "energyLevel = 'high', dailyBudget = maximumwaarde"],
            ["Elke grenswaarde", "Score op elke overgangswaarde", "Correct niveau en budget"],
            ["Deterministisch", "Zelfde score twee keer", "Altijd identiek resultaat"],
            ["Ongeldige invoer", "energyScore = 0, 11, -1, 'hoog'", "ZodError vóór berekening"],
        ],
    )

    p(doc, "9.2 Zod-schema's per domein", "Heading 2")
    table(
        doc,
        ["Schema", "Locatie", "Te testen gevallen"],
        [
            ["MorningCheckInSchema", "lib/checkin/schemas.ts", "Geldige check-in, score buiten bereik, te lange notitie"],
            ["OnboardingSubmissionSchema", "lib/onboarding/schemas.ts", "Geldige onboarding, ongeldige tijdzone"],
            ["SettingsSubmissionSchema", "lib/profile/schemas.ts", "Geldige settings, ongeldige herinneringstijd"],
            ["PlannedActivitySchema", "lib/planning/schemas.ts", "Geldige activiteit, negatieve energiepunten"],
        ],
    )

    p(doc, "9.3 Navigatie-utilities (risicoscore: Hoog)", "Heading 2")
    table(
        doc,
        ["Functie", "Bestand", "Te testen gevallen"],
        [
            ["sanitizeNextPath()", "lib/auth/navigation.ts", "Geldig pad, dubbele slash (//evil.com), externe URL, leeg pad"],
            ["buildPathWithQuery()", "lib/auth/navigation.ts", "Geen params, meerdere params, speciale tekens"],
            ["getAuthNotice()", "lib/auth/messages.ts", "Bekende foutcode, onbekende code, statuscode"],
        ],
    )

    # ── 10. Integratietests ──────────────────────────────────────────────────
    p(doc, "10. Integratietests", "Heading 1")
    p(
        doc,
        "Integratietests gebruiken een echte Supabase-testdatabase. Mocks worden niet ingezet voor de databaselaag: "
        "mock-gedrag verschilt van productiegedrag en verbergt RLS- en queryfouten.",
    )
    table(
        doc,
        ["Test", "Doel", "Risicoscore"],
        [
            ["getProfileBundleForCurrentUser()", "Retourneert gecombineerd profiel en settings", "Hoog"],
            ["ensureProfileBundleForCurrentUser()", "Maakt records aan als ze niet bestaan (bootstrap)", "Hoog"],
            ["createMorningCheckIn() — geldig", "Check-in opgeslagen, budget berekend en teruggegeven", "Kritiek"],
            ["createMorningCheckIn() — ongeldige invoer", "Zod fout vóór databaseschrijf", "Kritiek"],
            ["Rate limit: >10 snelle loginpogingen", "429-respons of vertraging aantoonbaar", "Hoog"],
            ["Reflectie-job idempotentie", "Dubbel uitvoeren geeft geen dubbele prompts", "Middel"],
        ],
    )

    # ── 11. RLS en security tests ────────────────────────────────────────────
    p(doc, "11. RLS en security tests (pgTAP)", "Heading 1")
    p(
        doc,
        "Elke tabel met gebruikersdata krijgt een eigen testbestand. "
        "Tests worden uitgevoerd als de 'authenticated'-rol met een gesimuleerd JWT. "
        "De SQL Editor-rol mag nooit worden gebruikt, omdat die RLS omzeilt.",
    )
    table(
        doc,
        ["Testgeval", "Te verifiëren"],
        [
            ["SELECT eigen rij", "Gebruiker A ziet alleen zijn eigen record"],
            ["SELECT andermans rij (IDOR)", "Gebruiker A ziet 0 rijen van gebruiker B"],
            ["INSERT voor zichzelf", "Aanmaken eigen record lukt"],
            ["INSERT voor een ander", "RLS-fout of 0 rows inserted"],
            ["UPDATE eigen rij", "Eigen record aanpassen lukt"],
            ["UPDATE andermans rij", "0 rows updated"],
            ["DELETE eigen rij", "Verwijderen eigen record lukt"],
            ["DELETE andermans rij", "0 rows deleted"],
            ["Unauthenticated (geen JWT)", "0 rijen, geen informatielekking in foutmelding"],
        ],
    )

    # ── 12. E2E tests ────────────────────────────────────────────────────────
    p(doc, "12. End-to-end tests (Playwright)", "Heading 1")
    table(
        doc,
        ["Flow", "Kritieke assertions", "Risicoscore"],
        [
            ["Registratie + e-mailbevestiging", "Dashboard bereikbaar na bevestiging", "Hoog"],
            ["Inloggen", "Dashboard zichtbaar, profiel aanwezig", "Hoog"],
            ["Beveiligde route zonder sessie", "Redirect naar /login", "Kritiek"],
            ["IDOR-poging (cross-user)", "Gebruiker B kan data van A niet ophalen", "Kritiek"],
            ["Ochtendcheck-in", "Budget en energieniveau zichtbaar na opslaan", "Kritiek"],
            ["Activiteit plannen + energiemeter", "Meter update direct, activiteit in overzicht", "Hoog"],
            ["Activiteit uitgevoerd / geskipt / aangepast", "Status correct in dagoverzicht", "Middel"],
            ["Instellingen wijzigen", "Succesbericht, instelling persistent na herlaad", "Middel"],
            ["Uitloggen", "Dashboard onbereikbaar na uitloggen", "Hoog"],
        ],
    )

    # ── 13. Testdata-management ──────────────────────────────────────────────
    p(doc, "13. Testdata-management", "Heading 1")
    bullets(
        doc,
        [
            "Elke E2E-test maakt eigen testdata aan of hergebruikt een dedicated testaccount. Nooit productiedata.",
            "Unit- en integratietests zijn stateless: geen gedeelde databaserecords.",
            "Gebruik zod-fixture om valide testobjecten te genereren vanuit Zod-schema's.",
            "Na integratietests worden records opgeruimd in afterEach of afterAll.",
            "Seed-scripts voor statische referentiedata staan in supabase/seed.sql.",
            "Gebruik een apart Supabase-testproject voor CI.",
        ],
    )

    # ── 14. CI/CD-integratie ─────────────────────────────────────────────────
    p(doc, "14. CI/CD-integratie", "Heading 1")
    table(
        doc,
        ["Job", "Trigger", "Stappen", "Blokkeerend voor merge"],
        [
            ["Lint en build", "PR en push naar main", "npm ci, npm run lint, npm run build", "Ja"],
            ["Unit en integratie", "PR en push naar main", "npm ci, npx vitest run, supabase test db", "Ja"],
            ["E2E", "PR naar main", "npm ci, npx playwright install, npx playwright test --workers=1", "Ja"],
            ["Security headers check", "PR naar main", "curl-check op staging-URL", "Ja"],
        ],
    )
    p(
        doc,
        "Mislukte tests blokkeren de merge. Overgeslagen tests vereisen expliciete goedkeuring "
        "inclusief gedocumenteerde risicoafweging (ISO 13485-principe: non-conformity met CAPA).",
    )

    # ── 15. Bestandsstructuur ────────────────────────────────────────────────
    p(doc, "15. Bestandsstructuur", "Heading 1")
    table(
        doc,
        ["Pad", "Inhoud"],
        [
            ["lib/checkin/__tests__/budget.test.ts", "Unit tests budgetberekening (risico: Kritiek)"],
            ["lib/checkin/schemas.ts", "Zod-schema MorningCheckIn"],
            ["lib/auth/__tests__/navigation.test.ts", "Unit tests open-redirect-preventie (risico: Hoog)"],
            ["lib/profile/__tests__/service.test.ts", "Integratietests profileservice"],
            ["supabase/tests/test_profiles_rls.sql", "pgTAP RLS-tests profiles"],
            ["supabase/tests/test_user_settings_rls.sql", "pgTAP RLS-tests user_settings"],
            ["supabase/tests/test_morning_check_ins_rls.sql", "pgTAP RLS-tests check-ins"],
            ["supabase/tests/test_activities_rls.sql", "pgTAP RLS-tests activiteiten"],
            ["e2e/auth.spec.ts", "Playwright: registratie, login, uitloggen, IDOR-poging"],
            ["e2e/checkin.spec.ts", "Playwright: ochtendcheck-in"],
            ["e2e/planning.spec.ts", "Playwright: activiteiten plannen en evalueren"],
            ["e2e/settings.spec.ts", "Playwright: instellingen"],
            ["playwright/global-setup.ts", "Programmatische Supabase-login, sessiestatus opslaan"],
            ["playwright.config.ts", "Playwright-configuratie incl. auth-setup en workers"],
            ["docs/traceability-matrix.md", "Levend document: FR-ID → test → resultaat per release"],
        ],
    )

    # ── 16. Acceptatiecriteria en Definition of Done ─────────────────────────
    p(doc, "16. Acceptatiecriteria en Definition of Done", "Heading 1")
    table(
        doc,
        ["Laag", "Minimale eis voor launch"],
        [
            ["Unit (Kritiek)", "Budgetberekening: 100% dekking op alle grenswaarden en ongeldige invoer."],
            ["Unit (Hoog)", "Zod-schema's getest op geldige en ongeldige invoer. Navigatie-utilities getest op open-redirect."],
            ["Integratie", "Profileservice: happy path en bootstrap. Check-in service: opslaan + budgetoutput. Rate limiting aantoonbaar actief."],
            ["RLS (pgTAP)", "Alle tabellen: SELECT/INSERT/UPDATE/DELETE als owner, als andere gebruiker en zonder sessie getest."],
            ["E2E", "Login, check-in, planning, evaluatie en uitloggen geautomatiseerd. IDOR-poging geeft geen data. Beveiligde route zonder sessie redirect."],
            ["Cybersecurity", "OWASP Top 10 scan zonder kritieke bevindingen. Security headers A-rating. AES-256 data at rest bevestigd."],
            ["Traceability (ISO 13485)", "Alle FR-IDs zijn gekoppeld aan een test. Kritieke FR's hebben gedocumenteerd testresultaat."],
            ["QMS", "Mislukte tests gedocumenteerd als non-conformity in Linear. Interne release-review van traceability matrix uitgevoerd."],
            ["Validatie (handmatig)", "Kernflows geverifieerd op mobiel. Usability test met minimaal 1 echte gebruiker. Copy getoetst op niet-medische formulering."],
        ],
    )

    # ── 17. Bewuste keuzes ───────────────────────────────────────────────────
    p(doc, "17. Bewuste keuzes en afwegingen", "Heading 1")
    table(
        doc,
        ["Keuze", "Alternatief", "Reden"],
        [
            ["Echte Supabase in integratietests", "Gemockte client", "Mocks verbergen RLS- en querygedrag. Mock/prod-divergentie is een bewezen risico."],
            ["pgTAP voor RLS", "Applicatielaag-tests", "RLS draait in de database; pgTAP test op hetzelfde uitvoerniveau."],
            ["Risicogebaseerde prioritering (ISO 14971)", "Gelijke dekking per module", "Testcapaciteit wordt ingezet waar de gevolgen van een fout het grootst zijn."],
            ["Traceability matrix (ISO 13485)", "Geen formele koppeling", "Vereiste voor auditeerbare kwaliteitsborging en MDR-voorbereiding."],
            ["NEN 7510 als toetssteen nu al", "Alleen OWASP", "Verlaagt drempel naar medische track, vermindert privacyrisico's bij gezondheidsdata."],
            ["ISO 13485 als leidraad (niet gecertificeerd)", "Geen QMS-structuur", "Bouwt documentregister en non-conformity-aanpak op zonder onnodige overhead."],
            ["Programmatische Playwright-login", "UI-login per test", "Sneller, minder foutgevoelig, scheidt auth-test van feature-test."],
        ],
    )

    # ── 18. Externe referenties ──────────────────────────────────────────────
    p(doc, "18. Externe referenties", "Heading 1")
    references = [
        ("Next.js Testing Guide — Vitest", "https://nextjs.org/docs/app/guides/testing/vitest"),
        ("Next.js Testing Guide — Playwright", "https://nextjs.org/docs/app/guides/testing/playwright"),
        ("Supabase Testing Overview", "https://supabase.com/docs/guides/local-development/testing/overview"),
        ("pgTAP documentatie", "https://pgtap.org/"),
        ("Zod documentatie", "https://zod.dev/"),
        ("zod-fixture — testdata genereren vanuit Zod-schema's", "https://github.com/timdeschryver/zod-fixture"),
        ("Playwright — Supabase auth via REST API", "https://mokkapps.de/blog/login-at-supabase-via-rest-api-in-playwright-e2e-test"),
        ("Playwright — opslaan en hergebruiken van auth-state", "https://playwright.dev/docs/auth"),
        ("ISO 14971 — Risicomanagement voor medische hulpmiddelen", "https://www.iso.org/standard/72704.html"),
        ("ISO 13485 — Kwaliteitsmanagementsystemen voor medische hulpmiddelen", "https://www.iso.org/standard/59752.html"),
        ("NEN 7510 — Informatiebeveiliging in de zorg", "https://www.nen.nl/nen-7510-1-2017-nl-237552"),
        ("OWASP Top 10 — Meest kritieke webapplicatierisico's", "https://owasp.org/www-project-top-ten/"),
        ("OWASP ZAP — Geautomatiseerde securityscanner", "https://www.zaproxy.org/"),
        ("SSL Labs — TLS-configuratiecheck", "https://www.ssllabs.com/ssltest/"),
        ("securityheaders.com — HTTP security headers checker", "https://securityheaders.com/"),
        ("EU MDR 2017/745 — Medical Device Regulation", "https://eur-lex.europa.eu/legal-content/EN/TXT/?uri=CELEX:32017R0745"),
    ]
    for name, url in references:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{name}: ")
        add_hyperlink(para, url, url)

    set_footer(doc, f"{PRODUCT_NAME} Testplan v0.2")
    doc.save(BASE_DIR / "inspannings-monitor-07-testplan-v01.docx")
    print("Testplan v0.2 gegenereerd.")


if __name__ == "__main__":
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    build_testplan()
