from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path("/Users/janpetervisser/Development/third/docs")
PRODUCT_NAME = "Inspannings Monitor"
DATE_TEXT = "18 april 2026"


def init_doc(title_text: str, subtitle_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos"
    styles["Title"].font.size = Pt(22)
    styles["Subtitle"].font.size = Pt(11)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(title_text)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(subtitle_text)
    doc.add_paragraph("")
    return doc


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Normal")


def label_value(doc: Document, label: str, value: str) -> None:
    para = doc.add_paragraph(style="Normal")
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    para.add_run(value)


PERSONAS = [
    {
        "name": "Marieke van den Berg",
        "age": 42,
        "gender": "Vrouw",
        "ethnicity": "Nederlands",
        "condition": "Burnout (herstelperiode)",
        "occupation": "Teamleider bij een gemeente",
        "location": "Apeldoorn",
        "story": (
            "Marieke groeide op in een gezin waar hard werken vanzelfsprekend was. Na haar "
            "studie Bestuurskunde in Nijmegen begon ze als beleidsmedewerker bij de gemeente "
            "Apeldoorn. Ze klom snel op. Op haar veertigste leidde ze een team van veertien "
            "mensen, was ze verantwoordelijk voor drie grote projecten en hielp ze haar moeder "
            "die net een heupoperatie had gehad.\n\n"
            "In november 2025 viel ze. Niet letterlijk, maar ze stond op een maandagochtend "
            "voor de deur van haar kantoor en kon niet naar binnen. Ze belde haar huisarts, "
            "die haar doorverwees naar een arbeidspsycholoog. Diagnose: ernstige burnout.\n\n"
            "Sindsdien is Marieke thuis. Ze heeft geleerd dat ze niet weet hoe ze naar haar "
            "eigen lichaam moet luisteren. Ze rijdt de dag vol met afspraken, ook als ze "
            "uitgeput is, en merkt het pas als ze op de bank zit en niet meer kan bewegen. "
            "Haar bedrijfsarts adviseerde haar een dagstructuur op basis van energiebeheer.\n\n"
            "Marieke is praktisch ingesteld. Ze wil geen dagboek bijhouden — ze wil een "
            "systeem dat haar helpt zonder dat ze er lang over na hoeft te denken. Ze gebruikt "
            "haar telefoon veel, maar wantrouwt apps die te veel vragen. Ze wil voortgang zien "
            "zonder dat het voelt als falen als ze een dag slecht scoort.\n\n"
            "Haar doel: over zes maanden weer twee dagen per week kunnen werken, zonder terug "
            "te vallen."
        ),
    },
    {
        "name": "Omar El Mansouri",
        "age": 35,
        "gender": "Man",
        "ethnicity": "Marokkaans-Nederlands",
        "condition": "Long COVID (vermoeidheid en cognitieve klachten)",
        "occupation": "Zelfstandig bouwondernemer",
        "location": "Rotterdam",
        "story": (
            "Omar is de oudste van vier kinderen. Zijn ouders kwamen in de jaren negentig uit "
            "Fez naar Rotterdam. Hij heeft zijn bedrijfje in renovatiewerk opgebouwd met zijn "
            "handen en zijn netwerk: weinig slapende nachten, veel werkende zondagen.\n\n"
            "In februari 2024 liep hij COVID op. De koorts duurde vijf dagen. Daarna ging het "
            "niet over. Na drie weken probeerde hij weer te werken. Na een uur op de steiger "
            "moest hij naar beneden. Zijn hoofd zat vol mist. Zijn handen trilden. Zijn borst "
            "trok samen als hij een trap opliep.\n\n"
            "Inmiddels weet hij dat hij long COVID heeft. Hij heeft een fysiotherapeut die hem "
            "op post-exertionele malaise heeft gewezen: als je te veel doet, crash je en doe je "
            "dagen terug. Zijn grootste probleem is dat hij dit niet voelt aankomen. Hij denkt "
            "dat hij het kan, totdat hij het niet meer kan.\n\n"
            "Omar wil zijn bedrijf niet verliezen. Hij heeft twee vaste medewerkers en een "
            "lopende hypotheek. Hij zoekt een manier om zijn dagen te plannen zodat hij niet "
            "elke woensdag knockdown is. Hij wil begrijpen wanneer hij meer en wanneer hij "
            "minder aankan. Een eenvoudige app op zijn telefoon, die snel in te vullen is en "
            "niet te veel vraagt, past bij hem."
        ),
    },
    {
        "name": "Priya Naipaul",
        "age": 51,
        "gender": "Vrouw",
        "ethnicity": "Surinaams-Nederlands",
        "condition": "Fibromyalgie",
        "occupation": "Leerkracht basisonderwijs (parttime)",
        "location": "Amsterdam Zuidoost",
        "story": (
            "Priya's ouders kwamen uit Suriname naar Nederland in 1980. Ze groeide op in "
            "Amsterdam Zuidoost, studeerde PABO en geeft al twintig jaar les op een "
            "basisschool om de hoek. Ze houdt van haar werk, maar de pijn maakt het steeds "
            "moeilijker.\n\n"
            "Op haar 46ste kreeg ze na jaren van onverklaarde pijn eindelijk een naam voor "
            "haar klachten: fibromyalgie. Haar lichaam reageert op stress, slaaptekort en "
            "overbelasting met hevige pijngolven die haar soms dagenlang in bed houden. Ze "
            "werkt nu drie dagen per week in plaats van vijf.\n\n"
            "Priya heeft geleerd dat haar energie per dag sterk wisselt. Op een goede dag "
            "kan ze boodschappen doen, een les geven en koken. Op een slechte dag is "
            "douchen al te veel. Ze wil beter voorspellen welke dag wat wordt, zodat ze "
            "niet steeds haar omgeving teleurstelt met last-minute afzeggingen.\n\n"
            "Ze is gewend aan apps en gebruikt haar telefoon voor alles. Maar ze heeft "
            "genoeg gehad van apps die haar het gevoel geven dat ze tekortschiet. Ze "
            "wil een tool die haar helpt plannen, niet één die haar beoordeelt. Haar "
            "man en twee volwassen kinderen steunen haar, maar begrijpen niet altijd "
            "waarom ze de ene dag 'normaal' doet en de andere dag volledig uitvalt."
        ),
    },
    {
        "name": "Jeroen Hoekstra",
        "age": 29,
        "gender": "Man",
        "ethnicity": "Nederlands",
        "condition": "ADHD en angststoornis",
        "occupation": "Junior softwareontwikkelaar",
        "location": "Utrecht",
        "story": (
            "Jeroen heeft ADHD die op zijn 24ste werd gediagnosticeerd, na jaren van "
            "uitstellen, vergeten en zichzelf opporren om te functioneren. Hij is slim "
            "en creatief, maar zijn energiehuishouding is chaotisch: hij kan uren "
            "hyperfocussen op een codetaak en vergeet dan te eten, of hij is zo "
            "overprikkeld na een dag vol meetings dat hij 's avonds niks meer kan.\n\n"
            "Naast zijn ADHD heeft hij een angststoornis die zich uit in piekeren: 's "
            "nachts, over zijn werk, over of hij genoeg doet, over of zijn collega's hem "
            "serieus nemen. De combinatie zorgt voor slaapproblemen die zijn overdag-"
            "functioneren verder ondermijnen.\n\n"
            "Hij woont alleen in een appartement in Utrecht, werkt hybride bij een "
            "softwarebedrijf en is in therapie bij een GZ-psycholoog. Zijn therapeut heeft "
            "hem geadviseerd zijn dag te structureren en momenten van herstel in te "
            "plannen. Jeroen wil dit doen, maar heeft hulp nodig om het vol te houden.\n\n"
            "Hij is techsavvy en heeft al tien apps geprobeerd. Hij stopt ermee als ze te "
            "ingewikkeld zijn, te veel notificaties sturen of te saai zijn. Hij wil een "
            "app die hem helpt de dag te starten met een helder beeld van wat er reëel "
            "is, en die hem niet straft als hij van plan verandert."
        ),
    },
    {
        "name": "Fatima Yilmaz",
        "age": 44,
        "gender": "Vrouw",
        "ethnicity": "Turks-Nederlands",
        "condition": "Multiple sclerose (relapsing-remitting)",
        "occupation": "Freelance vertaalster",
        "location": "Enschede",
        "story": (
            "Fatima is opgegroeid in Enschede als dochter van Turkse gastarbeiders. Ze "
            "studeerde Talen in Groningen en werkt nu als freelance vertaalster voor "
            "juridische teksten. Ze is moeder van twee kinderen van 12 en 15 jaar oud.\n\n"
            "Op haar 38ste kreeg ze de diagnose MS (relapsing-remitting). Sindsdien "
            "wisselen periodes van relatief goede gezondheid af met aanvallen waarbij "
            "ze last heeft van extreme vermoeidheid, tintelingen in haar handen en "
            "wazig zicht. Warmte maakt haar klachten altijd erger.\n\n"
            "Haar neuroloog heeft haar uitgelegd dat energie bij MS eindig is op een "
            "heel andere manier dan bij gezonde mensen. Ze heeft het concept 'the "
            "spoon theory' leren kennen en probeert er zo mee om te gaan. Maar ze "
            "heeft moeite om de grens aan te voelen voordat ze hem overschrijdt.\n\n"
            "Fatima werkt vanuit huis, wat flexibel is maar ook lastig: ze heeft geen "
            "vaste ritme en vraagdruk van klanten bepaalt haar dag. Ze wil een systeem "
            "dat haar helpt de balans te vinden tussen haar werk, haar gezin en haar "
            "gezondheid. Ze is goed met computers maar wil geen extra schermen: "
            "telefoon is makkelijker, zeker als haar handen tintelen."
        ),
    },
    {
        "name": "Kenneth Asare",
        "age": 58,
        "gender": "Man",
        "ethnicity": "Ghanees-Nederlands",
        "condition": "Diabetes type 2 met vermoeidheidsklachten",
        "occupation": "Chauffeur (zit momenteel in ziekteverlof)",
        "location": "Den Haag",
        "story": (
            "Kenneth kwam op zijn 22ste vanuit Ghana naar Nederland om te studeren, maar "
            "belandde in de logistiek. Hij werkte dertig jaar als chauffeur voor een "
            "transportbedrijf en was trots op zijn betrouwbaarheid: hij was zelden ziek.\n\n"
            "Vier jaar geleden werd bij hem diabetes type 2 vastgesteld. Hij paste zijn "
            "eetpatroon aan, nam zijn medicatie, maar de vermoeidheid bleef. Zijn "
            "huisarts legde hem uit dat schommelende bloedsuikerspiegels zijn energie "
            "direct beïnvloeden: soms voelt hij zich prima, soms zit hij na het "
            "ontbijt al leeg. Na een incident waarbij hij bijna de controle over zijn "
            "vrachtwagen verloor door hypoglycemie, is hij met verlof gestuurd.\n\n"
            "Kenneth is niet gewend aan stilzitten. Hij heeft moeite met de "
            "onzekerheid: welke dagen zijn de slechte? Wanneer kan hij iets plannen "
            "met zijn kleinkinderen? Hij wil grip terugkrijgen op zijn dag. Hij is "
            "minder gewend aan smartphones dan jongere gebruikers maar leert snel als "
            "de interface duidelijk is. Zijn vrouw helpt hem soms, maar hij wil zelf "
            "in controle zijn.\n\n"
            "Zijn doel is om te begrijpen wat zijn energiepatronen zijn, zodat hij "
            "straks — als hij wellicht deeltijd kan hervatten — weet wanneer hij wat "
            "aankan."
        ),
    },
    {
        "name": "Lisa Vermeulen",
        "age": 23,
        "gender": "Vrouw",
        "ethnicity": "Nederlands",
        "condition": "ME/CVS (chronisch vermoeidheidssyndroom)",
        "occupation": "Voormalig student (studie gestaakt)",
        "location": "Breda",
        "story": (
            "Lisa studeerde Psychologie in Tilburg toen ze ziek werd, twee jaar geleden. "
            "Wat begon als een virale infectie, werd nooit beter. Na negen maanden "
            "van steeds teruggaan naar de huisarts met klachten die ze niet goed kon "
            "verwoorden, werd ME/CVS vastgesteld.\n\n"
            "Ze heeft haar studie moeten staken. Ze woont weer bij haar ouders in Breda "
            "omdat ze niet voor zichzelf kan zorgen op drukke dagen. Op goede dagen "
            "maakt ze een kleine wandeling en kan ze een uur achter haar laptop zitten. "
            "Op slechte dagen is ze bedlegerig en kan ze het licht en geluid in haar "
            "kamer nauwelijks verdragen.\n\n"
            "Lisa's grootste uitdaging is de onzichtbaarheid van haar ziekte. Ze ziet "
            "er uit als een gezonde 23-jarige. Vrienden begrijpen niet waarom ze niet "
            "naar feestjes kan. Ze heeft geleerd dat post-exertionele malaise haar "
            "grootste vijand is: te veel doen kost haar dagenlang herstel.\n\n"
            "Ze wil grip krijgen op haar beschikbare energie. Niet om meer te doen, "
            "maar om beter te verdelen wat ze heeft. Ze is vertrouwd met technologie "
            "en heeft haar telefoon vlakbij, maar heeft apps nodig die snel te "
            "gebruiken zijn — lange formulieren zijn op slechte dagen onmogelijk. "
            "Ze wil ook haar eigen data kunnen terugzien, zodat ze het gesprek met "
            "haar behandelaars beter kan voeren."
        ),
    },
    {
        "name": "Sander Tan",
        "age": 47,
        "gender": "Man",
        "ethnicity": "Indonesisch-Nederlands",
        "condition": "Depressie (recidiverend) en slaapstoornis",
        "occupation": "Grafisch ontwerper (deeltijd)",
        "location": "Haarlem",
        "story": (
            "Sander is de kleinzoon van Indische Nederlanders die na de dekolonisatie "
            "naar Nederland kwamen. Hij groeide op in Haarlem in een gezin dat weinig "
            "sprak over emoties of gezondheid. Hij werkte jarenlang als grafisch "
            "ontwerper voor reclamebureaus, creatief en productief, totdat zijn eerste "
            "depressie hem op zijn 38ste volledig lam legde.\n\n"
            "Sindsdien heeft hij drie depressieve episodes doorgemaakt. Tussen de "
            "episodes door werkt hij parttime, maar hij voelt zich nooit helemaal "
            "stabiel. Hij slaapt slecht: soms te weinig, soms te lang, zelden "
            "herstellend. De slaaptekort versterkt zijn stemming negatief en omgekeerd.\n\n"
            "Zijn psychiater heeft hem geleerd dat bij depressie de energiehuishouding "
            "verstoord is: zijn motivatie en zijn vermogen zijn niet altijd in balans. "
            "Hij heeft moeite om te beginnen aan dingen, zelfs als hij weet dat ze "
            "hem goed doen. Maar als hij begint, kan hij soms te lang doorgaan en dan "
            "crasht hij de volgende dag.\n\n"
            "Sander zoekt een lichte structuur voor zijn dag. Niet te dwingend — dat "
            "verhoogt zijn angst — maar genoeg houvast om niet de hele dag te "
            "improviseren. Hij is visueel ingesteld en waardeert een heldere, "
            "rustige interface. Hij wil zien hoe hij het gedaan heeft, maar niet "
            "worden overspoeld met grafieken. Eén duidelijk getal is genoeg."
        ),
    },
]


FACE_IMAGES = [f"/tmp/persona_face_{i}.jpg" for i in range(1, 9)]


def add_photo(doc: Document, image_path: str) -> None:
    from pathlib import Path as _Path
    if not _Path(image_path).exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    run.add_picture(image_path, width=Inches(1.5))


def build_personas():
    doc = init_doc(
        f"{PRODUCT_NAME} — Gebruikerspersonas",
        f"Fictieve gebruikers voor ontwerp en testvalidatie · {DATE_TEXT}",
    )

    doc.add_heading("Inleiding", level=1)
    p(doc, (
        "Dit document beschrijft acht fictieve gebruikers van de Inspannings Monitor. "
        "Ze zijn ontworpen om de diversiteit van de doelgroep te vertegenwoordigen: "
        "volwassenen die te maken hebben met energiebeheer als gevolg van uiteenlopende "
        "gezondheidsklachten. De personas zijn fictief maar gebaseerd op reële ziektebeelden "
        "en ervaringen die relevant zijn voor de positionering van het product als "
        "wellness/self-management tool."
    ))
    p(doc, (
        "De personas zijn verdeeld naar geslacht (4 vrouwen, 4 mannen), etniciteit "
        "(Nederlandse achtergrond, Marokkaans-Nederlands, Turks-Nederlands, Surinaams-"
        "Nederlands, Ghanees-Nederlands, Indonesisch-Nederlands) en ziektebeeld "
        "(burnout, long COVID, fibromyalgie, ADHD/angst, MS, diabetes type 2, ME/CVS, "
        "depressie/slaapstoornis)."
    ))
    p(doc, (
        "Gebruik: de personas dienen als referentiekader voor UX-beslissingen, "
        "testscenario's en prioritering van functies. Ze zijn geen medische profielen "
        "en vormen geen basis voor medische claims."
    ))

    doc.add_heading("Overzicht", level=1)
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Naam", "Geslacht / leeftijd", "Achtergrond", "Ziektebeeld"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for i, persona in enumerate(PERSONAS, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = persona["name"]
        row[2].text = f"{persona['gender']}, {persona['age']}"
        row[3].text = persona["ethnicity"]
        row[4].text = persona["condition"]

    doc.add_paragraph("")

    doc.add_heading("Gebruikerspersonas", level=1)

    for i, persona in enumerate(PERSONAS, start=1):
        doc.add_heading(f"{i}. {persona['name']}", level=2)

        add_photo(doc, FACE_IMAGES[i - 1])

        label_value(doc, "Leeftijd", str(persona["age"]))
        label_value(doc, "Geslacht", persona["gender"])
        label_value(doc, "Etniciteit", persona["ethnicity"])
        label_value(doc, "Beroep", persona["occupation"])
        label_value(doc, "Woonplaats", persona["location"])
        label_value(doc, "Ziektebeeld", persona["condition"])

        doc.add_heading("Levensverhaal", level=3)
        for paragraph in persona["story"].split("\n\n"):
            p(doc, paragraph.strip())

        doc.add_paragraph("")

    out_path = BASE_DIR / "inspannings-monitor-08-gebruikerspersonas-v01.docx"
    doc.save(out_path)
    print(f"Opgeslagen: {out_path}")


if __name__ == "__main__":
    build_personas()
