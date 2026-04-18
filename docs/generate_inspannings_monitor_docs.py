from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt


BASE_DIR = Path("/Users/janpetervisser/Development/third/docs")
PRODUCT_NAME = "Inspannings Monitor"
DATE_TEXT = "17 april 2026"
POSITIONING = "Wellness/self-management"
LANGUAGES = "Nederlands"
HOSTING = "Vercel"
DATABASE = "Supabase PostgreSQL"
AUTH = "Supabase Auth"
AUDIENCE = "Volwassenen"


def init_doc(title_text: str, subtitle_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos"
    styles["Title"].font.size = Pt(22)
    styles["Subtitle"].font.size = Pt(11)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(title_text)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(subtitle_text)
    doc.add_paragraph("")
    return doc


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)
    run.append(run_props)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def p(doc: Document, text: str = "", style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def bullets(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, headers, rows) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph("")


def set_footer(doc: Document, text: str) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = text


def build_productkader() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Productkader en Positionering v0.6",
        f"{POSITIONING}-route met expliciet opengehouden future medical track\n{DATE_TEXT}",
    )
    p(doc, "1. Documentdoel", "Heading 1")
    p(
        doc,
        f"Dit document legt de strategische productpositie van {PRODUCT_NAME} vast voor de eerstvolgende versies. "
        "Het voorkomt dat marketing, product, ontwerp en engineering verschillende aannames gebruiken over wat het product wel en niet is. "
        "De gekozen route is bewust: eerst een wellness/self-management product, later eventueel een aparte medische producttrack als daar bewijs, middelen en governance voor zijn.",
    )

    p(doc, "2. Besluit", "Heading 1")
    p(
        doc,
        f"{PRODUCT_NAME} wordt in de MVP en de direct daaropvolgende releases gepositioneerd als een wellness/self-management product. "
        "Het ondersteunt zelfobservatie, planning en reflectie rondom energie, belasting en herstel. "
        "Het product is in deze fase uitdrukkelijk geen medisch hulpmiddel, geen diagnostisch systeem en geen behandelondersteunend klinisch systeem.",
    )

    p(doc, "3. Intended Use", "Heading 1")
    p(
        doc,
        f"{PRODUCT_NAME} helpt {AUDIENCE.lower()} gebruikers met wisselende of beperkte energiereserves om hun dagelijkse activiteiten te plannen, "
        "uit te voeren en te evalueren in relatie tot hun ervaren energie. Het product ondersteunt gebruikers bij zelfmanagement door een eenvoudige plan-doe-evalueer cyclus, "
        "wekelijkse terugblik en optionele reflectieprompts na zwaardere dagen.",
    )

    p(doc, "4. Non-Intended Use", "Heading 1")
    bullets(
        doc,
        [
            f"{PRODUCT_NAME} is niet bedoeld voor diagnose, behandeling, preventie, monitoring of voorspelling van een ziekte in medische zin.",
            f"{PRODUCT_NAME} geeft geen patiëntspecifiek medisch advies, geen triage en geen therapeutische aanbevelingen.",
            f"{PRODUCT_NAME} is niet bedoeld om behandelbesluiten van een arts of andere zorgprofessional te sturen.",
            f"{PRODUCT_NAME} is niet bedoeld als noodhulpmiddel of voor spoedsituaties.",
            f"{PRODUCT_NAME} is niet bedoeld als formele bron voor opname in een patiëntendossier in de wellness-first fase.",
        ],
    )

    p(doc, "5. Productprincipes", "Heading 1")
    bullets(
        doc,
        [
            "Minimale gebruikslast: de gebruiker heeft weinig energie; elke interactie moet zuinig zijn.",
            "Geen oordeel: het product toont patronen, geen morele duiding.",
            "Privacy by default: alleen noodzakelijke data, standaard zo beperkt mogelijk.",
            "Uitlegbare inzichten: transparante regels boven slimme maar ondoorzichtige uitkomsten.",
            "Wellness claims alleen: copy en UI mogen de gekozen intended use niet overschrijden.",
            "Medical-ready fundament: documentatie, logging en architectuur moeten latere aanscherping mogelijk maken zonder nu al medische claims te maken.",
        ],
    )

    p(doc, "6. Primaire doelgroep", "Heading 1")
    bullets(
        doc,
        [
            f"{AUDIENCE} met chronische of terugkerende vermoeidheid die meer grip willen op hun dagindeling en energieverdeling.",
            "Mensen die baat hebben bij gestructureerde zelfobservatie zonder behoefte aan een klinische app of zorgintegratie in de eerste fase.",
        ],
    )

    p(doc, "7. Rollen in de wellness-first fase", "Heading 1")
    table(
        doc,
        ["Rol", "Toegelaten in MVP", "Toelichting"],
        [
            ["Gebruiker/eigenaar", "Ja", "Beheert eigen gegevens, check-ins, activiteiten en persoonlijke instellingen."],
            ["Zorgverlener", "Nee", "Geen viewerrol in release 1; blijft buiten scope tot een latere fase."],
            ["Naaste/mantelzorger", "Nee", "Geen live delen in release 1; privacy en veiligheidskaders eerst aanscherpen."],
            ["Admin/support", "Beperkt", "Alleen systeembeheer waar strikt noodzakelijk, met logging en rolgebonden toegang."],
        ],
    )

    p(doc, "8. MVP-scope", "Heading 1")
    bullets(
        doc,
        [
            "Ochtend energie check-in met vaste schaal en slaapkwaliteit.",
            "Activiteiten plannen voor de dag binnen een eenvoudig energiebudgetmodel.",
            "Activiteiten markeren als uitgevoerd, geskipt of aangepast.",
            "Dagoverzicht en weekoverzicht met eenvoudige, uitlegbare patronen.",
            "Optionele reflectieprompt op T+1 en T+2 na een zware of overschreden dag.",
            "Basisinstellingen voor taal, timezone, zichtbaarheid van energiebudgetpunten en herinneringen.",
        ],
    )

    p(doc, "9. Uitdrukkelijk buiten scope voor MVP", "Heading 1")
    bullets(
        doc,
        [
            "Delen met zorgverleners, naasten of andere viewers.",
            "PDF-export voor patiëntendossiers of formele zorgrapportage.",
            "Medicatie-tracking, drugs/middelen-logging, nicotine-logging en andere extra gevoelige habit-tracking.",
            "AI-gegenereerde inzichten of vrije tekstinterpretatie door een model.",
            "Claims over voorspellen van PEM, ziekteverloop of behandeladvies.",
            "Integraties met EPD, wearables of externe medische databronnen.",
        ],
    )

    p(doc, "10. Claim- en copy-guardrails", "Heading 1")
    bullets(
        doc,
        [
            "Wel toegestaan: “helpt je plannen”, “maakt patronen zichtbaar”, “ondersteunt zelfmanagement”, “helpt je terugkijken op je energieverdeling”.",
            "Niet toegestaan: “voorspelt crashes”, “detecteert PEM”, “ondersteunt behandelbesluiten”, “geschikt voor patiëntendossier”, “klinisch gevalideerd” tenzij dat later aantoonbaar klopt.",
            "Elke insightkaart gebruikt patroon-taal: “in jouw data lijkt … samen te hangen met …”, nooit causale taal.",
            "Elke reflectieprompt verwijst niet naar risico of medische noodzaak, maar naar zelfreflectie en terugblik.",
        ],
    )

    p(doc, "11. Bevestigde productkeuzes voor deze versie", "Heading 1")
    table(
        doc,
        ["Onderwerp", "Besluit"],
        [
            ["Productnaam", PRODUCT_NAME],
            ["Doelgroep", AUDIENCE],
            ["Release 1", "Alleen individuele gebruikers"],
            ["Voertaal eerste release", LANGUAGES],
            ["Hosting", HOSTING],
            ["Database", DATABASE],
            ["Authenticatie", AUTH],
        ],
    )

    p(doc, "12. Voorwaarden om later een medische producttrack te starten", "Heading 1")
    numbered(
        doc,
        [
            "Er is een expliciet strategisch besluit om intended use te verbreden naar een medisch doel.",
            "Er is aparte documentatie voor de medische variant of medische release-tak.",
            "Claims, clinician workflows en rapportage worden opnieuw beoordeeld op MDR-consequenties.",
            "Er is budget en eigenaarschap voor risicomanagement, klinische evaluatie, kwaliteitsmanagement en post-market verplichtingen.",
            "De wellness-versie en de toekomstige medische variant krijgen duidelijke change control en traceerbaarheid.",
        ],
    )

    p(doc, "13. Succescriteria voor de wellness-first fase", "Heading 1")
    bullets(
        doc,
        [
            "Gebruikers kunnen de kernloop dagelijks gebruiken zonder hoge cognitieve belasting.",
            "Gebruikers ervaren de app als ondersteunend en niet normatief of veroordelend.",
            "De eerste release verwerkt alleen data die nodig is voor het beoogde wellness-doel.",
            "De documentatie is sterk genoeg voor productreview, privacyreview en securityreview.",
            "Het product blijft inhoudelijk en qua copy binnen de wellness-positionering.",
        ],
    )

    set_footer(doc, f"{PRODUCT_NAME} Productkader en Positionering v0.6")
    doc.save(BASE_DIR / "inspannings-monitor-01-productkader-en-positionering-v06.docx")


def build_functionele_specificatie() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Functionele Specificatie MVP v0.6",
        f"{POSITIONING}-scope met toetsbare requirements\n{DATE_TEXT}",
    )
    p(doc, "1. Documentdoel", "Heading 1")
    p(
        doc,
        "Dit document beschrijft de functionele scope van de wellness-first MVP. Het is nadrukkelijk geen roadmapdocument en geen privacydocument; "
        "alleen gedrag, regels, staten en acceptance criteria die nodig zijn om het product te bouwen en te testen.",
    )

    p(doc, "2. Kerngebruikersreis", "Heading 1")
    numbered(
        doc,
        [
            "De gebruiker opent de app en doet een korte ochtendcheck-in.",
            "De app bepaalt een eenvoudig dagbudget op basis van de gekozen energieschaal.",
            "De gebruiker plant enkele activiteiten voor de dag.",
            "De app toont een visuele meter en een niet-blokkerende waarschuwing bij overschrijding.",
            "Gedurende of na de dag markeert de gebruiker activiteiten als uitgevoerd, geskipt of aangepast.",
            "De gebruiker bekijkt het dagoverzicht en later het weekoverzicht om patronen te zien.",
            "Na een zware dag kan de app een optionele reflectieprompt tonen op T+1 en T+2.",
        ],
    )

    p(doc, "3. Functionele modules", "Heading 1")
    p(doc, "3.1 Onboarding", "Heading 2")
    bullets(
        doc,
        [
            "Korte onboarding van maximaal drie schermen.",
            "Uitleg over energieschaal, dagplanning en niet-medische positionering.",
            "Keuze van taal, timezone en herinneringsinstellingen.",
            "Geen uitgebreide profielvragen of gevoelige aanvullende intake in MVP.",
        ],
    )

    p(doc, "3.2 Ochtendcheck-in", "Heading 2")
    bullets(
        doc,
        [
            "Energiescore via visuele schaal met ankerlabels.",
            "Slaapkwaliteit als vast veld.",
            "Optionele korte notitie.",
            "Automatische afleiding van energieniveau en dagbudget.",
        ],
    )

    p(doc, "3.3 Activiteiten plannen", "Heading 2")
    bullets(
        doc,
        [
            "Velden: naam, categorie, geplande duur, energie-impact, optioneel tijdslot, prioriteit.",
            "Autocomplete op basis van eerder ingevoerde activiteiten.",
            "Lopend totaal van geplande energiebelasting tegenover dagbudget.",
            "Niet-blokkerende feedback als het geplande totaal het budget overschrijdt.",
        ],
    )

    p(doc, "3.4 Uitvoeren en evalueren", "Heading 2")
    bullets(
        doc,
        [
            "Statussen: uitgevoerd, geskipt, aangepast.",
            "Bij uitgevoerd: werkelijke duur, vermoeidheidsscore na afloop, optionele notitie.",
            "Bij geskipt: reden uit lijst plus optionele toelichting.",
            "Bij aangepast: omschrijving van aanpassing, werkelijke duur, vermoeidheidsscore.",
            "Ongeplande activiteiten mogen worden toegevoegd met dezelfde basisvelden.",
        ],
    )

    p(doc, "3.5 Overzichten", "Heading 2")
    bullets(
        doc,
        [
            "Dagoverzicht met geplande versus uitgevoerde activiteiten, totale energiebelasting en korte samenvatting.",
            "Weekoverzicht met eenvoudige patronen, adherence aan budget en skip-patronen.",
            "Inzichten blijven transparant, beperkt en uitlegbaar.",
        ],
    )

    p(doc, "3.6 Reflectie na zware dagen", "Heading 2")
    bullets(
        doc,
        [
            "Optionele T+1 en T+2 reflectie na dagen met duidelijke overschrijding of hoge belasting.",
            "Reflectie bestaat uit korte zelfrapportage, niet uit medische waarschuwing.",
            "Prompt is opt-in en kan door gebruiker worden uitgezet.",
        ],
    )

    p(doc, "4. Functionele requirements", "Heading 1")
    requirements = [
        ["FR-ONB-001", "Onboarding", "De onboarding bestaat uit maximaal drie schermen en kan worden overgeslagen of later worden afgerond.", "Acceptatie: gebruiker kan onboarding afronden binnen circa 1 minuut zonder verplichte gevoelige invoer."],
        ["FR-CHK-001", "Check-in", "De gebruiker kan een ochtendcheck-in opslaan met energiescore en slaapkwaliteit.", "Acceptatie: opslaan lukt met beide velden ingevuld; andere velden zijn optioneel."],
        ["FR-CHK-002", "Check-in", "Het systeem leidt uit de energiescore automatisch een energieniveau en dagbudget af.", "Acceptatie: dagbudget wordt direct zichtbaar na opslaan van de check-in."],
        ["FR-PLAN-001", "Planning", "De gebruiker kan een activiteit plannen met naam, categorie, duur, energie-impact en prioriteit.", "Acceptatie: activiteit verschijnt direct in het dagoverzicht."],
        ["FR-PLAN-002", "Planning", "De app toont steeds het geplande totaal versus dagbudget.", "Acceptatie: totaal wordt bijgewerkt na elke mutatie."],
        ["FR-PLAN-003", "Planning", "Bij overschrijding toont de app een niet-blokkerende waarschuwing.", "Acceptatie: gebruiker kan bewust doorgaan zonder blokkade."],
        ["FR-ACT-001", "Uitvoering", "Een geplande activiteit kan worden gemarkeerd als uitgevoerd.", "Acceptatie: werkelijke duur en vermoeidheidsscore zijn invulbaar en worden opgeslagen."],
        ["FR-ACT-002", "Uitvoering", "Een geplande activiteit kan worden gemarkeerd als geskipt.", "Acceptatie: skip-reden kan gekozen worden uit een lijst met optionele toelichting."],
        ["FR-ACT-003", "Uitvoering", "Een geplande activiteit kan worden gemarkeerd als aangepast.", "Acceptatie: aanpassingstekst, werkelijke duur en vermoeidheidsscore kunnen worden vastgelegd."],
        ["FR-ACT-004", "Uitvoering", "De status “aangepast” dekt ook deels uitgevoerde activiteiten.", "Acceptatie: geen aparte status nodig om dit scenario vast te leggen."],
        ["FR-ACT-005", "Uitvoering", "De gebruiker kan een ongeplande activiteit toevoegen.", "Acceptatie: ongeplande activiteit telt mee in het werkelijke totaal."],
        ["FR-DAY-001", "Dagoverzicht", "De app toont het verschil tussen gepland en uitgevoerd.", "Acceptatie: scherm bevat beide totalen en de relevante activiteitenstatussen."],
        ["FR-DAY-002", "Dagoverzicht", "De gebruiker ziet een korte samenvatting van de dag zonder medische interpretatie.", "Acceptatie: tekst gebruikt patroon- of reflectietaal, geen medische taal."],
        ["FR-WEEK-001", "Weekoverzicht", "De app toont per week gemiddelde energie, budget-adherence en skip-patronen.", "Acceptatie: gebruiker kan een week terugkijken zonder ruwe database-informatie te hoeven interpreteren."],
        ["FR-INS-001", "Inzichten", "Een inzicht wordt alleen getoond als aan minimale datadrempels is voldaan.", "Acceptatie: bij te weinig data toont de app geen patroonclaim maar een neutrale melding."],
        ["FR-REM-001", "Reflectie", "De gebruiker kan optionele reflectieprompts ontvangen na een zware dag.", "Acceptatie: prompts zijn per gebruiker aan of uit te zetten."],
        ["FR-SET-001", "Instellingen", "De gebruiker kan taal, timezone, herinneringen en zichtbaarheid van punten beheren.", "Acceptatie: wijzigingen zijn direct actief op het account."],
        ["FR-ACC-001", "Toegankelijkheid", "Belangrijke acties zijn mobiel bruikbaar met grote touch targets.", "Acceptatie: primaire knoppen voldoen aan de gekozen UX-richtlijn voor mobiel gebruik."],
    ]
    table(doc, ["ID", "Module", "Requirement", "Acceptance / toetsing"], requirements)

    p(doc, "5. Bevestigde platformkeuzes voor deze MVP", "Heading 1")
    table(
        doc,
        ["Onderwerp", "Besluit", "Effect op scope"],
        [
            ["Productnaam", PRODUCT_NAME, "Alle UI-copy en documentatie gebruiken deze naam."],
            ["Doelgroep", AUDIENCE, "Geen flows voor minderjarigen in de eerste release."],
            ["Taal", LANGUAGES, "Geen meertaligheidsvereisten in release 1."],
            ["Authenticatie", AUTH, "Account- en sessiebeheer volgt de Supabase-architectuur."],
            ["Database", DATABASE, "Datamodel en beveiliging worden rond Supabase/PostgreSQL ontworpen."],
            ["Hosting", HOSTING, "Webapp-hosting en deployment gaan uit van Vercel."],
        ],
    )

    p(doc, "6. Datavelden in MVP", "Heading 1")
    table(
        doc,
        ["Domein", "Velden in MVP", "Opmerkingen"],
        [
            ["Profiel", "naam of schermnaam, taal, timezone", "Zo minimaal mogelijk."],
            ["Ochtendcheck-in", "energiescore, slaapkwaliteit, optionele notitie, timestamp", "Geen uitgebreid symptomenformulier in MVP."],
            ["Activiteit", "naam, categorie, geplande duur, energie-impact, prioriteit, tijdslot optioneel, status, werkelijke duur, fatigue na afloop, optionele notitie", "Kern van het product."],
            ["Weekpatronen", "berekende totalen en eenvoudige aggregaties", "Alleen uitlegbare regels."],
            ["Reflectieprompt", "korte follow-up score en optionele notitie", "Opt-in en niet medisch geformuleerd."],
        ],
    )

    p(doc, "7. Expliciet buiten scope van deze specificatie", "Heading 1")
    bullets(
        doc,
        [
            "Caregiver- of professional-sharing.",
            "Habit tracking buiten slaapkwaliteit en kernenergieflow.",
            "Medicatie, alcohol, nicotine, drugs/middelen, water- of voedingstracking.",
            "AI-samenvattingen, generatieve tekstuitleg en chatbotfuncties.",
            "PDF-export of andere zorgdossier-achtige rapportages.",
            "Integraties met wearable data of medische systemen.",
        ],
    )

    p(doc, "8. Release-acceptatie voor MVP", "Heading 1")
    numbered(
        doc,
        [
            "Alle kernflows werken op mobiel zonder hoge interactielast.",
            "De app blijft binnen de wellness/non-medical copy-guardrails.",
            "Geen viewerrollen of deelroutes zijn actief in release 1.",
            "Inzichten gebruiken alleen expliciet gedefinieerde regels.",
            "Alle requirements met FR-ID zijn aantoonbaar getest of handmatig geverifieerd.",
        ],
    )

    set_footer(doc, f"{PRODUCT_NAME} Functionele Specificatie MVP v0.6")
    doc.save(BASE_DIR / "inspannings-monitor-02-functionele-specificatie-mvp-v06.docx")


def build_privacy_security_safety() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Privacy, Security en Safety Baseline v0.2",
        f"Basiseisen voor een wellness-first product met gezondheidsgerelateerde data\n{DATE_TEXT}",
    )
    p(doc, "1. Documentdoel", "Heading 1")
    p(
        doc,
        "Dit document beschrijft de minimale privacy-, beveiligings- en safety-baseline voor de eerste wellness-first releases. "
        "Het is opgesteld om te voorkomen dat gevoelige datafunctionality sneller groeit dan governance, beveiliging en risicoafdekking.",
    )

    p(doc, "2. Werkhypothese voor gegevensverwerking", "Heading 1")
    p(
        doc,
        f"{PRODUCT_NAME} verwerkt gegevens die in context gezondheidsgerelateerd en vaak bijzonder gevoelig zijn. Daarom wordt in dit document uitgegaan van een strikte benadering: "
        "dataminimalisatie, expliciete toestemming waar nodig, terughoudendheid met extra datacategorieen en een DPIA voorafgaand aan echte gebruikersintroductie. "
        "Juridische validatie blijft nodig, maar de productdocumentatie moet nu al vanuit de strengere lezing worden ontworpen.",
    )

    p(doc, "3. Bevestigde technische uitgangspunten", "Heading 1")
    table(
        doc,
        ["Onderwerp", "Bevestigd uitgangspunt", "Implicatie"],
        [
            ["Hosting", HOSTING, "Deployment- en security-instellingen moeten aansluiten op het Vercel-model."],
            ["Database", DATABASE, "Gegevensmodellering, toegangscontrole en back-upstrategie steunen op Supabase/PostgreSQL."],
            ["Authenticatie", AUTH, "Identiteit, sessies en autorisatie gaan uit van Supabase Auth."],
            ["Releasevorm", "Alleen individuele gebruikers", "Geen sharing-architectuur in release 1."],
            ["Taal", LANGUAGES, "Geen meertalige opslag- of vertaalpijplijn in launchscope."],
        ],
    )

    p(doc, "4. Dataminimalisatie en default-instellingen", "Heading 1")
    table(
        doc,
        ["Datadomein", "Toelaatbaar in MVP", "Default", "Opmerking"],
        [
            ["Account/profiel", "Ja", "Minimaal", "Alleen wat nodig is voor account, taal en timezone."],
            ["Energiescore en slaapkwaliteit", "Ja", "Actief", "Kerngegevens voor intended use."],
            ["Activiteitenplanning", "Ja", "Actief", "Kerngegevens voor intended use."],
            ["Vrije notities", "Beperkt", "Optioneel", "Korte notities toestaan, geen uitgebreide journaling als kernflow."],
            ["Reflectie na zware dag", "Ja", "Opt-in", "Alleen als wellness-reflectieprompt."],
            ["Delen met derden", "Nee", "Uit", "Niet opnemen in release 1."],
            ["Medicatie", "Nee", "Uit", "Te gevoelig voor wellness-first MVP."],
            ["Alcohol/nicotine/drugs", "Nee", "Uit", "Niet noodzakelijk voor eerste value proof."],
            ["Geavanceerde leefstijlfactoren", "Nee", "Uit", "Alleen later na aparte risicoafweging."],
        ],
    )

    p(doc, "5. Aanbevolen privacy-eisen", "Heading 1")
    table(
        doc,
        ["ID", "Eis"],
        [
            ["PRIV-001", "Voor elke datacategorie moet een expliciet doel in documentatie zijn vastgelegd."],
            ["PRIV-002", "Niet-noodzakelijke datavelden staan standaard uit of bestaan nog niet in de MVP."],
            ["PRIV-003", "De gebruiker kan eigen data exporteren en verwijderen via een duidelijk proces."],
            ["PRIV-004", "Bewaartermijnen worden per datadomein vastgesteld voor launch."],
            ["PRIV-005", "Er is vóór launch een DPIA uitgevoerd op de daadwerkelijke MVP-scope."],
            ["PRIV-006", "Subverwerkers, regio’s en dataflows zijn in een apart register vastgelegd."],
            ["PRIV-007", "Marketingcopy en onboarding mogen geen bredere verwerking suggereren dan feitelijk plaatsvindt."],
        ],
    )

    p(doc, "6. Aanbevolen security-eisen", "Heading 1")
    table(
        doc,
        ["ID", "Eis"],
        [
            ["SEC-001", "Alle netwerkcommunicatie verloopt via moderne TLS-configuratie."],
            ["SEC-002", "Gegevens worden versleuteld opgeslagen op platformniveau en waar nodig aanvullend logisch afgescheiden."],
            ["SEC-003", "Sessiebeheer voorkomt onbeperkte of onzichtbare langdurige toegang."],
            ["SEC-004", "Rate limiting en basisbescherming tegen misbruik zijn actief op auth- en mutatieroutes."],
            ["SEC-005", "Beheer- en supporttoegang is rolgebonden, minimaal en auditbaar."],
            ["SEC-006", "Belangrijke gebeurtenissen zoals login, mislukte toegang, data-export en accountverwijdering worden gelogd."],
            ["SEC-007", "Back-ups en hersteltests zijn gedefinieerd vóór launch."],
            ["SEC-008", "Secrets worden niet in code of clientside configuratie opgeslagen."],
        ],
    )

    p(doc, "7. Safety- en content guardrails", "Heading 1")
    table(
        doc,
        ["ID", "Eis"],
        [
            ["SAFE-001", "Het product gebruikt geen diagnose-, behandel- of voorspellende taal in de wellness-first fase."],
            ["SAFE-002", "Reflectieprompts worden gepresenteerd als zelfreflectie, niet als medische waarschuwing."],
            ["SAFE-003", "Inzichten tonen alleen patronen bij voldoende data en vermelden periode en aantal observaties waar relevant."],
            ["SAFE-004", "De app toont geen advies om medicatie te wijzigen, zorg uit te stellen of medisch handelen te vervangen."],
            ["SAFE-005", "De onboarding en helpteksten bevatten een heldere noodgevallen- en niet-medische disclaimer."],
            ["SAFE-006", "Nieuwe claims of clinician-facing functies vereisen aparte review voordat ze in product en copy verschijnen."],
        ],
    )

    p(doc, "8. Hoogste productrisico’s in deze fase", "Heading 1")
    table(
        doc,
        ["Risico", "Waarom relevant", "Mitigatie in deze documentatie"],
        [
            ["Te brede datascope", "Meer gevoelige data dan nodig verhoogt risico en governance-last.", "Scope beperken tot energie, slaapkwaliteit, activiteiten en korte reflectie."],
            ["Te medische copy", "Kan product regulatoir opschuiven en gebruikers verkeerd laten vertrouwen.", "Intended use en non-intended use expliciet vastleggen."],
            ["Schijnzekerheid in inzichten", "Gebruiker kan patronen als feiten of voorspellingen lezen.", "Alleen eenvoudige, uitlegbare regels en guardrails."],
            ["Onvoldoende zicht op datatoegang", "Gezondheidsgerelateerde data vragen auditability.", "Audit logging en rolgebonden beheer als launch-voorwaarde."],
            ["Premature sharing", "Delen met derden vergroot privacy- en securityrisico sterk.", "Geen sharing in release 1; pas later na apart ontwerp."],
        ],
    )

    p(doc, "9. Pre-launch artefacten die verplicht gereed moeten zijn", "Heading 1")
    numbered(
        doc,
        [
            "DPIA op de werkelijke MVP-scope.",
            "Datacatalogus met per veld doel, gevoeligheid, bewaartermijn en toegangsmodel.",
            "Subprocessor- en regio-overzicht voor onder meer Vercel en Supabase.",
            "Beveiligingsbaselines voor authenticatie, logging, back-up en incidentafhandeling.",
            "Safety review van productcopy, onboarding en insightformuleringen.",
            "Besluitdocument waarin expliciet is vastgelegd dat delen met derden buiten de launchscope valt.",
        ],
    )

    p(doc, "10. Externe referenties", "Heading 1")
    references = [
        ("EDPB Data Protection Basics", "https://www.edpb.europa.eu/sme-data-protection-guide/data-protection-basics_en"),
        ("EDPB Be Compliant", "https://www.edpb.europa.eu/sme-data-protection-guide/be-compliant_en"),
        ("GDPR via EUR-Lex", "https://eur-lex.europa.eu/eli/reg/2016/679/oj/eng"),
        ("NHS DTAC", "https://transform.england.nhs.uk/key-tools-and-info/digital-technology-assessment-criteria-dtac/"),
        ("NHS Clinical Risk Management Standards", "https://digital.nhs.uk/services/clinical-safety/clinical-risk-management-standards"),
        ("HL7 FHIR Security & Privacy", "https://www.hl7.org/fhir/secpriv-module.html"),
        ("European Commission MDCG guidance index", "https://health.ec.europa.eu/medical-devices-sector/new-regulations/guidance-mdcg-endorsed-documents-and-other-guidance_en"),
        ("Nictiz over NEN 7510", "https://nationalebibliotheek.nictiz.nl/bibliotheek/nen-7510/"),
    ]
    for name, url in references:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"{name}: ")
        add_hyperlink(para, url, url)

    set_footer(doc, f"{PRODUCT_NAME} Privacy, Security en Safety Baseline v0.2")
    doc.save(BASE_DIR / "inspannings-monitor-03-privacy-security-safety-baseline-v02.docx")


def build_roadmap() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Roadmap: Wellness-first naar Eventuele Medische Track v0.2",
        f"Gefaseerde productontwikkeling met expliciete decision gates\n{DATE_TEXT}",
    )
    p(doc, "1. Doel van deze roadmap", "Heading 1")
    p(
        doc,
        f"Deze roadmap beschrijft hoe {PRODUCT_NAME} gecontroleerd kan groeien vanuit een wellness/self-management MVP naar een mogelijk zwaardere productvorm in de toekomst, "
        "zonder nu al medische claims of medische scope binnen te halen. De roadmap is bewust gefaseerd om eerst productwaarde, governance en veiligheid op orde te brengen.",
    )

    p(doc, "2. Fase 0: Fundament vóór launch", "Heading 1")
    bullets(
        doc,
        [
            "Definitieve intended use en non-intended use goedkeuren.",
            "MVP-scope vastleggen en extra gevoelige datacategorieen expliciet uitsluiten.",
            "DPIA uitvoeren op de werkelijke launchscope.",
            f"{AUTH}, {DATABASE} en {HOSTING} als startarchitectuur formaliseren.",
            "Productcopy, onboarding en inzichtteksten safety-reviewen.",
            "Support- en incidentproces benoemen.",
        ],
    )

    p(doc, "3. Fase 1: Wellness-first MVP", "Heading 1")
    bullets(
        doc,
        [
            "Ochtendcheck-in, dagplanning, uitvoering/evaluatie, dag- en weekoverzicht.",
            "Optionele reflectieprompt na zware dagen.",
            "Geen delen met derden, geen habit-uitbreiding, geen AI, geen zorgrapportages.",
            "Voertaal bij launch: alleen Nederlands.",
            "Hoofddoel: bewijzen dat de kernloop waarde oplevert en laagdrempelig is.",
        ],
    )

    p(doc, "4. Fase 2: Beperkte uitbreiding binnen wellness", "Heading 1")
    bullets(
        doc,
        [
            "Verbeterde weekpatronen en transparante insightregels.",
            "Eventueel beperkte extra gewoontevelden met lage gevoeligheid, alleen na aparte review.",
            "Sjablonen en gebruiksgemakverbeteringen.",
            "Doel: betere retentie en reflectiewaarde zonder positionering te verbreden.",
        ],
    )

    p(doc, "5. Fase 3: Gecontroleerd delen buiten MVP", "Heading 1")
    bullets(
        doc,
        [
            "Pas starten na apart ontwerp voor consent, viewer-authenticatie, audit logging en toegangsbeheer.",
            "Begin bij voorkeur met een smalle share-scope en geen vrije notities of hooggevoelige velden.",
            "Geen “zorgdashboard” taal zolang intended use wellness blijft.",
            "Doel: delen onderzoeken zonder onbedoeld naar klinische besluitondersteuning te schuiven.",
        ],
    )

    p(doc, "6. Future Medical Track: aparte producttak, geen kleine feature-toggle", "Heading 1")
    p(
        doc,
        f"Als {PRODUCT_NAME} later een medisch product of medisch geclassificeerde software wil worden, dan moet dat worden behandeld als een aparte producttrack of minimaal als een expliciet gereguleerde release-tak. "
        "Het mag niet ontstaan doordat losse features langzaam over de grens schuiven.",
    )
    numbered(
        doc,
        [
            "Start met een formeel besluit over nieuwe intended use.",
            "Herbeoordeel claims, clinician workflows, rapportages en data-interpretatie op MDR-impact.",
            "Richt aparte documentatie in voor risicomanagement, klinische evaluatie en change control.",
            "Bepaal welke onderdelen van de wellness-architectuur herbruikbaar zijn en welke opnieuw ontworpen moeten worden.",
            "Maak een duidelijke scheiding tussen de wellness-productlijn en de medische productlijn in communicatie en releasebeheer.",
        ],
    )

    p(doc, "7. Decision gates", "Heading 1")
    table(
        doc,
        ["Gate", "Vraag", "Minimale voorwaarde om door te gaan"],
        [
            ["Gate A", "Is de wellness-positionering intern en extern scherp genoeg?", "Approved intended use, non-intended use en copy-guardrails."],
            ["Gate B", "Is launch juridisch en organisatorisch verantwoord?", "DPIA, datacatalogus, security baseline en supportproces gereed."],
            ["Gate C", "Mag sharing worden toegevoegd?", "Aparte consent-, auth- en audit-architectuur plus risicoreview."],
            ["Gate D", "Is een medical track gewenst?", "Expliciet business- en productbesluit met budget en ownership."],
            ["Gate E", "Mag een medical track worden gestart?", "Nieuwe intended use, regulatoire analyse en extra kwaliteitsartefacten zijn ingepland."],
        ],
    )

    p(doc, "8. Signalering dat het product richting medisch schuift", "Heading 1")
    bullets(
        doc,
        [
            "Er verschijnen claims over voorspellen, detecteren of verminderen van ziekte of symptomen.",
            "Zorgverleners gaan op de output vertrouwen voor behandelafstemming.",
            "PDF- of dashboardfunctionaliteit wordt ontworpen voor formele zorgdossiers.",
            "Het systeem geeft patiëntspecifieke aanbevelingen die verder gaan dan algemene zelfreflectie.",
            "Marketing of sales gebruikt klinische taal om waarde te beschrijven.",
        ],
    )

    p(doc, "9. Bevestigde uitgangspunten voor de eerstvolgende maanden", "Heading 1")
    table(
        doc,
        ["Onderwerp", "Bevestigd uitgangspunt"],
        [
            ["Productnaam", PRODUCT_NAME],
            ["Doelgroep", AUDIENCE],
            ["Voertaal launch", LANGUAGES],
            ["Launchvorm", "Alleen individuele gebruikers"],
            ["Hosting", HOSTING],
            ["Authenticatie", AUTH],
            ["Database", DATABASE],
        ],
    )

    p(doc, "10. Aanbevolen werkvolgorde voor de komende weken", "Heading 1")
    numbered(
        doc,
        [
            "Maak deze v0.6-documentenset leidend en archiveer verouderde aannames uit eerdere documenten.",
            "Werk de MVP uit in backlog-items op basis van requirement-ID’s.",
            "Plan DPIA, security baseline en copy review parallel aan ontwerp en bouw.",
            "Formuleer Vercel- en Supabase-keuzes als formeel architectuurbesluit met subprocessor-overzicht.",
            "Gebruik elke scope-uitbreiding langs de decision gates in dit document.",
        ],
    )

    set_footer(doc, f"{PRODUCT_NAME} Roadmap: Wellness-first naar Eventuele Medische Track v0.2")
    doc.save(BASE_DIR / "inspannings-monitor-04-roadmap-wellness-naar-medisch-v02.docx")


def build_technische_architectuur() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Technische Architectuur en Implementatie v0.1",
        f"Technische uitwerking van de wellness-first MVP op {HOSTING} en {AUTH}\n{DATE_TEXT}",
    )
    p(doc, "1. Documentdoel", "Heading 1")
    p(
        doc,
        "Dit document beschrijft de technische implementatielaag van de wellness-first MVP. "
        "Het vertaalt de product- en privacykeuzes naar een concrete architectuur, datamodel, autorisatiemodel, deploy-aanpak en teststrategie. "
        "Dit document is bewust apart gehouden van de functionele specificatie, zodat producteisen en technische keuzes niet opnieuw door elkaar gaan lopen.",
    )

    p(doc, "2. Bevestigde technische uitgangspunten", "Heading 1")
    table(
        doc,
        ["Onderwerp", "Keuze", "Toelichting"],
        [
            ["Frontend hosting", HOSTING, "De webapp wordt in eerste instantie gedeployed op Vercel."],
            ["Database", DATABASE, "De primaire datastore is Supabase PostgreSQL."],
            ["Authenticatie", AUTH, "Identity, sessies en basisautorisatie lopen via Supabase Auth."],
            ["Applicatietype", "Webapp, mobile-first", "Eerste release richt zich op individueel gebruik op mobiel, met bruikbare desktopweergave."],
            ["Release 1", "Alleen individuele gebruikers", "Geen share-dashboard, viewerrollen of zorgverlenerstoegang in de MVP."],
            ["Voertaal", LANGUAGES, "Geen meertaligheidsinfrastructuur in de launchscope."],
        ],
    )

    p(doc, "3. Doelarchitectuur op hoofdlijnen", "Heading 1")
    bullets(
        doc,
        [
            "Client: browsergebaseerde UI, primair mobiel, responsief voor desktop.",
            "App-laag: Next.js App Router met server-side rendering waar dat de dataflow en veiligheid vereenvoudigt.",
            "Auth-laag: Supabase Auth voor accounts, sessies en gebruikersidentiteit.",
            "Data-laag: Supabase PostgreSQL voor kernentiteiten en geaggregeerde overzichten.",
            "Business logic: server actions en server-side services voor mutaties, validatie en berekeningen.",
            "Presentatielaag: eenvoudige, uitlegbare inzichten op basis van transparante regels, niet op basis van AI.",
        ],
    )

    p(doc, "4. Aanbevolen stackindeling", "Heading 1")
    table(
        doc,
        ["Laag", "Aanbevolen keuze", "Reden"],
        [
            ["Framework", "Next.js met App Router", "Past goed bij Vercel, server components en server-side datatoegang."],
            ["UI", "React + TypeScript", "Sterke typeveiligheid en goede componentstructuur."],
            ["Styling", "Tailwind CSS + componentbibliotheek", "Snel consistente, toegankelijke UI voor mobile-first flows."],
            ["Validatie", "Zod of vergelijkbare runtime-validatie", "Nodig voor strikte server-side inputcontrole."],
            ["Database access", "Supabase client en/of Prisma", "Prisma kan nog steeds zinvol zijn, maar is een bewuste keuze en geen verplichting."],
            ["Charts", "Lichte grafiekbibliotheek", "Alleen voor eenvoudige dag- en weekinzichten; geen heavy analytics platform nodig."],
            ["Background jobs", "Cron via Vercel of vergelijkbare jobtrigger", "Voldoende voor simpele reminder- en onderhoudstaken in MVP."],
        ],
    )

    p(doc, "5. Systeemcontext", "Heading 1")
    table(
        doc,
        ["Component", "Verantwoordelijkheid"],
        [
            ["Browserclient", "Rendert UI, verzamelt gebruikersinvoer, toont overzichten en roept beveiligde mutaties aan."],
            ["Next.js serverlaag", "Rendert pagina’s, bewaakt routes, valideert mutaties, voert domeinlogica uit."],
            ["Supabase Auth", "Gebruikersaccounts, sessies, identity en auth-context."],
            ["Supabase Postgres", "Opslag van profieldata, check-ins, activiteiten, instellingen en geaggregeerde weekdata."],
            ["Cron/reminderlaag", "Plant en verstuurt optionele reflectie- en onderhoudstaken."],
            ["Logging/monitoring", "Verzamelt fouten, audit-events en operationele signalen."],
        ],
    )

    p(doc, "6. Applicatiestructuur", "Heading 1")
    p(doc, "6.1 Routevoorstel", "Heading 2")
    table(
        doc,
        ["Route", "Type", "Doel"],
        [
            ["/", "Server", "Landing en instappunt naar login of dashboard."],
            ["/login", "Server + Client", "Authenticatieflow via Supabase Auth."],
            ["/onboarding", "Client", "Korte uitleg en eerste instellingen."],
            ["/dashboard", "Server + Client", "Dagoverzicht met check-instatus, energiemeter en activiteiten."],
            ["/plan", "Client", "Activiteiten plannen of aanpassen voor de dag."],
            ["/history", "Server", "Weekoverzicht en eenvoudige patronen."],
            ["/history/[date]", "Server", "Dagdetail en terugblik op een specifieke dag."],
            ["/settings", "Client", "Taal, timezone, reminders, zichtbaarheid van punten en accountacties."],
        ],
    )

    p(doc, "6.2 Kerncomponenten", "Heading 2")
    bullets(
        doc,
        [
            "EnergySlider voor de ochtendscore.",
            "SleepQualityInput als vast onderdeel van de ochtendcheck-in.",
            "EnergyMeter voor budget versus gepland of werkelijk verbruik.",
            "ActivityCard voor geplande en geëvalueerde activiteiten.",
            "QuickCheckIn voor snelle handelingen met lage interactielast.",
            "DaySummary voor geplande versus uitgevoerde totalen.",
            "WeekTrendChart voor eenvoudige wekelijkse patronen.",
        ],
    )

    p(doc, "7. Authenticatie en autorisatie", "Heading 1")
    bullets(
        doc,
        [
            f"Authenticatie loopt via {AUTH}; de applicatie gebruikt de auth-context server-side om de ingelogde gebruiker te bepalen.",
            "Release 1 kent alleen het concept eigenaar/gebruiker; er zijn geen viewerrollen.",
            "Elke mutatie en elke data-opvraag moet server-side gekoppeld worden aan de ingelogde eigenaar.",
            "Beheer- of supporttoegang moet gescheiden blijven van gewone gebruikersaccounts en auditbaar zijn.",
            "Routes met persoonlijke data worden alleen geserveerd wanneer een geldige sessie aanwezig is.",
        ],
    )

    p(doc, "8. Row-Level Security (RLS) strategie", "Heading 1")
    p(
        doc,
        "Hoewel sharing niet in de MVP zit, is RLS nog steeds waardevol. De basisregel in release 1 is eenvoudig: een gebruiker mag uitsluitend records lezen en muteren die aan diens eigen profiel gekoppeld zijn. "
        "Dat geeft een verdedigbare securitybasis en voorkomt dat autorisatie alleen in applicatiecode leeft.",
    )
    table(
        doc,
        ["Tabelgroep", "RLS-principe in MVP"],
        [
            ["Profile / UserSettings", "Alleen owner mag eigen profiel en instellingen lezen en wijzigen."],
            ["EnergyCheckIn", "Alleen owner mag eigen check-ins lezen en schrijven."],
            ["PlannedActivity", "Alleen owner mag eigen activiteiten lezen en muteren."],
            ["SkipReason / ActivityCategory (user-defined)", "Alleen owner mag eigen uitbreidingen zien en beheren."],
            ["Aggregaties / weekdata", "Alleen owner mag afgeleide weekinzichten opvragen."],
        ],
    )

    p(doc, "9. Datamodel voor de wellness-first MVP", "Heading 1")
    table(
        doc,
        ["Model", "Kernvelden", "Opmerking"],
        [
            ["Profile", "id, displayName, locale, timezone, onboardingCompleted", "Profiel gekoppeld aan auth identity."],
            ["UserSettings", "profileId, morningReminder, reflectionReminderEnabled, showEnergyPoints", "Alleen settings die in MVP nodig zijn."],
            ["EnergyCheckIn", "id, profileId, score, energyLevel, dailyBudget, sleepQuality, note, timestamp", "Ochtendcheck-in en budgetbasis."],
            ["ActivityCategory", "id, profileId nullable, name, type, isSystem, sortOrder", "Systeemcategorieën plus optionele eigen categorieën."],
            ["PlannedActivity", "id, profileId, name, categoryId, plannedDate, plannedTimeSlot, plannedDurationMin, energyImpact, priority, status, actualDurationMin, fatigueScoreAfter, skipReasonId, adjustmentNote, note, isUnplanned, createdAt", "Kernentiteit voor plan/doe/evalueer."],
            ["SkipReason", "id, profileId nullable, name, isSystem, isActive, sortOrder", "Systeemredenen plus optionele eigen redenen."],
            ["ReflectionCheckIn", "id, profileId, relatedDate, checkInDate, fatigueScore, note", "Wellness-reflectie op T+1/T+2; niet medisch formuleren."],
        ],
    )

    p(doc, "10. Dataflow van de kernloop", "Heading 1")
    numbered(
        doc,
        [
            "Gebruiker authenticereert en opent het dashboard.",
            "Server haalt profiel, settings, laatste check-in en dagactiviteiten op voor de ingelogde gebruiker.",
            "Gebruiker doet een ochtendcheck-in; server valideert invoer en berekent energyLevel en dailyBudget.",
            "Gebruiker plant activiteiten; server valideert, schrijft records weg en berekent het lopend totaal.",
            "Gebruiker markeert activiteiten als uitgevoerd, geskipt of aangepast; server slaat evaluatievelden op.",
            "Dag- en weekoverzichten lezen geaggregeerde waarden uit of berekenen die server-side.",
            "Een reminderjob kan optioneel bepalen of een reflectieprompt op T+1 of T+2 klaarstaat.",
        ],
    )

    p(doc, "11. Server actions en services", "Heading 1")
    table(
        doc,
        ["Actie/service", "Doel"],
        [
            ["createMorningCheckIn", "Slaat energiescore en slaapkwaliteit op en berekent dagbudget."],
            ["planActivity", "Voegt een geplande activiteit toe aan een dag."],
            ["updatePlannedActivity", "Wijzigt duur, impact, tijdslot of prioriteit van een activiteit."],
            ["completeActivity", "Markeert activiteit als uitgevoerd met werkelijke duur en fatigue na afloop."],
            ["skipActivity", "Markeert activiteit als geskipt met reden en toelichting."],
            ["adjustActivity", "Legt aangepaste/deels uitgevoerde activiteit vast."],
            ["createUnplannedActivity", "Voegt een ongeplande activiteit toe aan de dag."],
            ["getDailyOverview", "Levert alle gegevens voor het dashboard van een specifieke dag."],
            ["getWeeklyOverview", "Levert geaggregeerde weekinzichten en patronen."],
            ["saveSettings", "Wijzigt taal, timezone en reminderinstellingen."],
            ["createReflectionCheckIn", "Slaat optionele T+1/T+2 reflectie op."],
        ],
    )

    p(doc, "12. Inzichten-engine", "Heading 1")
    bullets(
        doc,
        [
            "In release 1 gebruikt de app alleen regelgebaseerde, uitlegbare inzichten.",
            "Elke insightregel moet een minimale datadrempel kennen; zonder voldoende data wordt geen patroonclaim getoond.",
            "Inzichten worden server-side berekend om clientcomplexiteit en inconsistentie te beperken.",
            "Taal blijft patroon-georiënteerd en niet-medisch.",
            "AI en vrije tekstuitleg door modellen blijven expliciet buiten scope van deze technische versie.",
        ],
    )

    p(doc, "13. Scheduling en reminders", "Heading 1")
    bullets(
        doc,
        [
            "Release 1 vereist alleen eenvoudige geplande taken voor reflectieprompts en eventueel basisopschoning.",
            "Een cron-trigger op Vercel of vergelijkbare scheduled mechanismen is voldoende.",
            "Jobs moeten idempotent zijn, zodat dubbel uitvoeren niet tot dubbele prompts of dubbele records leidt.",
            "Elke reminder blijft opt-in en gekoppeld aan gebruikersinstellingen.",
        ],
    )

    p(doc, "14. Deployment en omgevingsscheiding", "Heading 1")
    table(
        doc,
        ["Onderdeel", "Aanpak"],
        [
            ["Omgevingen", "Minimaal development, preview/staging en production."],
            ["Hosting", "Next.js-app via Vercel deployments per omgeving."],
            ["Database", "Supabase-projecten of logisch gescheiden omgevingen per fase."],
            ["Secrets", "Opslaan in omgevingsbeheer van Vercel en Supabase, nooit in clientbundles of repository."],
            ["Migrations", "Schemawijzigingen via gecontroleerde migraties, niet handmatig in productie."],
            ["Rollbacks", "Deploystrategie moet snelle rollback van frontend mogelijk maken; databasemigraties krijgen een apart rollbackplan."],
        ],
    )

    p(doc, "15. Logging, monitoring en audit", "Heading 1")
    bullets(
        doc,
        [
            "Applicatiefouten en server exceptions worden centraal gelogd.",
            "Belangrijke security-events zoals loginfouten, sessieproblemen en accountverwijdering worden apart vastgelegd.",
            "Mutaties op check-ins, activiteiten en settings moeten herleidbaar zijn op technisch niveau.",
            "Ondersteuningsacties van beheerders worden apart auditbaar gemaakt.",
            "Monitoring moet minimaal inzicht geven in availability, cronfouten en foutpercentages van kernmutaties.",
        ],
    )

    p(doc, "16. Teststrategie", "Heading 1")
    table(
        doc,
        ["Testlaag", "Wat moet worden afgedekt"],
        [
            ["Unit tests", "Budgetberekening, insightregels, mapping van energiescore naar level en validatielogica."],
            ["Integration tests", "Server actions, auth-context, data-opslag en RLS-gerelateerde toegangspaden."],
            ["UI tests", "Ochtendcheck-in, activiteitenflow, dagoverzicht en instellingen op mobiel formaat."],
            ["Policy tests", "Controle dat gebruikers alleen hun eigen data kunnen lezen en muteren."],
            ["Manual QA", "Toegankelijkheid, lage-interactielast, foutmeldingen en regressie in kernflows."],
        ],
    )

    p(doc, "17. Bewuste afwijkingen ten opzichte van v0.4", "Heading 1")
    bullets(
        doc,
        [
            "Geen deelmodule, viewerrollen of shared dashboards in de MVP-architectuur.",
            "Geen gewoonte- of middelenmodule buiten slaapkwaliteit in release 1.",
            "Geen database-gestuurde meertaligheid in de launchfase; alleen Nederlands is bevestigd.",
            "Geen AI-laag of medische copy in de technische scope.",
            "PDF-export, zorgdossierdoelen en zorgverlenerstoegang zijn expliciet uit deze eerste architectuur gehaald.",
        ],
    )

    p(doc, "18. Vooruitkijkend ontwerp", "Heading 1")
    p(
        doc,
        "De architectuur mag future-ready zijn, maar niet over-abstract. Dat betekent: RLS en duidelijke domeinmodellen nu al goed neerzetten, "
        "maar sharing, meertaligheid, geavanceerde habits of medische varianten pas toevoegen wanneer daar een apart besluit en aparte specificatie voor is.",
    )

    set_footer(doc, f"{PRODUCT_NAME} Technische Architectuur en Implementatie v0.1")
    doc.save(BASE_DIR / "inspannings-monitor-05-technische-architectuur-en-implementatie-v01.docx")


def build_implementatieplan_backlog() -> None:
    doc = init_doc(
        f"{PRODUCT_NAME} Implementatieplan en Backlog v0.1",
        f"Vertaling van product-, privacy- en architectuurkeuzes naar epics en werkpakketten\n{DATE_TEXT}",
    )
    p(doc, "1. Documentdoel", "Heading 1")
    p(
        doc,
        "Dit document vertaalt de huidige documentatieset naar een uitvoerbaar implementatieplan voor release 1. "
        "Het groepeert werk in epics en backlogblokken, benoemt afhankelijkheden en maakt duidelijk wat eerst gebouwd moet worden en wat bewust later blijft.",
    )

    p(doc, "2. Uitgangspunten", "Heading 1")
    bullets(
        doc,
        [
            f"Productnaam: {PRODUCT_NAME}.",
            f"Positionering: {POSITIONING}.",
            "Release 1 is alleen voor individuele gebruikers.",
            f"Doelgroep: {AUDIENCE.lower()}.",
            f"Voertaal in release 1: {LANGUAGES}.",
            f"Technische basis: {HOSTING} + {AUTH} + {DATABASE}.",
            "Geen sharing, geen AI, geen PDF-export, geen medische claims in release 1.",
        ],
    )

    p(doc, "3. Aanbevolen bouwvolgorde", "Heading 1")
    numbered(
        doc,
        [
            "Fundament en projectopzet.",
            "Authenticatie, profiel en instellingen.",
            "Ochtendcheck-in en budgetlogica.",
            "Activiteiten plannen.",
            "Activiteiten evalueren en dagoverzicht.",
            "Weekoverzicht en uitlegbare inzichten.",
            "Reflectieprompts en geplande taken.",
            "Privacy, security, logging en launch-readiness.",
        ],
    )

    p(doc, "4. Epic-overzicht", "Heading 1")
    table(
        doc,
        ["Epic", "Doel", "Prioriteit", "Afhankelijk van"],
        [
            ["EPIC-01 Fundament", "Projectbasis, CI, omgevingen en design foundation neerzetten.", "P0", "-"],
            ["EPIC-02 Auth en profiel", "Inloggen, sessies, profiel en basisinstellingen werkend maken.", "P0", "EPIC-01"],
            ["EPIC-03 Ochtendcheck-in", "Energiescore, slaapkwaliteit en dagbudget implementeren.", "P0", "EPIC-02"],
            ["EPIC-04 Dagplanning", "Activiteiten plannen en budgetfeedback tonen.", "P0", "EPIC-03"],
            ["EPIC-05 Evaluatie en dagoverzicht", "Activiteiten afronden en dagresultaat tonen.", "P0", "EPIC-04"],
            ["EPIC-06 Weekoverzicht en inzichten", "Weekpatronen en veilige insightregels toevoegen.", "P1", "EPIC-05"],
            ["EPIC-07 Reflectie en reminders", "Optionele T+1/T+2 follow-up mogelijk maken.", "P1", "EPIC-05"],
            ["EPIC-08 Security en operations", "Logging, auditability, back-up, rate limiting en hardening.", "P0", "EPIC-01 t/m EPIC-07"],
            ["EPIC-09 Launch-readiness", "QA, content review, DPIA-input en go-live checks afronden.", "P0", "EPIC-01 t/m EPIC-08"],
        ],
    )

    p(doc, "5. Epic 01: Fundament", "Heading 1")
    p(doc, "Doel: een stabiele technische basis waarop alle kernflows kunnen landen.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-001", "Next.js projectbasis opzetten met TypeScript en gekozen stylingaanpak.", "Build", "Project start lokaal en in previewomgeving zonder handmatige workarounds."],
            ["ST-002", "Omgevingen definiëren voor development, preview en production.", "Ops", "Environment strategy is vastgelegd en werkt technisch."],
            ["ST-003", "Component foundation voor formulieren, kaarten, knoppen en meldingen neerzetten.", "UI", "Kerncomponenten zijn herbruikbaar en mobiel bruikbaar."],
            ["ST-004", "Basale foutafhandeling en lege staten ontwerpen.", "UX", "Gebruiker ziet bruikbare feedback bij lege of foutieve situaties."],
        ],
    )

    p(doc, "6. Epic 02: Auth en profiel", "Heading 1")
    p(doc, "Doel: iedere gebruiker kan veilig een eigen account en basisinstellingen beheren.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-101", "Supabase Auth integreren in de app en sessieflow inrichten.", "Build", "Gebruiker kan inloggen en beveiligde routes gebruiken."],
            ["ST-102", "Profile- en UserSettings-model implementeren.", "Build", "Profiel- en settingsrecords zijn per gebruiker beschikbaar."],
            ["ST-103", "Onboardingflow van maximaal drie schermen implementeren.", "UX", "Nieuwe gebruiker begrijpt schaal, positionering en basisinstellingen."],
            ["ST-104", "Settingsscherm bouwen voor taal, timezone, reminders en zichtbaarheid van punten.", "Build", "Wijzigingen worden persistent opgeslagen."],
            ["ST-105", "RLS-basispolicies voor owner-only toegang inrichten.", "Security", "Gebruiker kan uitsluitend eigen profiel en settings lezen of wijzigen."],
        ],
    )

    p(doc, "7. Epic 03: Ochtendcheck-in", "Heading 1")
    p(doc, "Doel: de gebruiker kan met minimale inspanning de dag starten en een budget krijgen.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-201", "EnergySlider en SleepQualityInput component bouwen.", "UI", "Check-in kan mobiel comfortabel worden ingevuld."],
            ["ST-202", "Server action voor createMorningCheckIn implementeren.", "Build", "Check-in wordt opgeslagen met juiste validatie."],
            ["ST-203", "Logica voor mapping van score naar energyLevel en dailyBudget bouwen.", "Logic", "Budget wordt consistent en testbaar berekend."],
            ["ST-204", "Check-instatus en budget direct zichtbaar maken op dashboard.", "UI", "Gebruiker ziet onmiddellijk het resultaat van de check-in."],
            ["ST-205", "Unit tests voor score- en budgetmapping toevoegen.", "QA", "Belangrijkste grenswaarden zijn afgedekt."],
        ],
    )

    p(doc, "8. Epic 04: Dagplanning", "Heading 1")
    p(doc, "Doel: de gebruiker kan activiteiten voor de dag plannen binnen een eenvoudig energiemodel.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-301", "Datamodel voor activiteiten, categorieën en skip-redenen implementeren.", "Build", "Migraties en basisseed-data zijn aanwezig."],
            ["ST-302", "Planningformulier bouwen met naam, categorie, duur, impact en prioriteit.", "UI", "Gebruiker kan een activiteit aanmaken zonder onnodige complexiteit."],
            ["ST-303", "Autocomplete op basis van eerdere activiteiten toevoegen.", "UX", "Veelgebruikte activiteiten zijn snel opnieuw te kiezen."],
            ["ST-304", "EnergyMeter en lopend totaal implementeren.", "Logic/UI", "Totaal update direct na elke wijziging."],
            ["ST-305", "Niet-blokkerende waarschuwing bij budgetoverschrijding toevoegen.", "UX", "Gebruiker krijgt feedback maar behoudt regie."],
        ],
    )

    p(doc, "9. Epic 05: Evaluatie en dagoverzicht", "Heading 1")
    p(doc, "Doel: de kernloop afronden door geplande activiteiten te evalueren en terug te zien.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-401", "Statusflows voor uitgevoerd, geskipt en aangepast implementeren.", "Build", "Alle drie de statussen kunnen correct worden opgeslagen."],
            ["ST-402", "Formuliervelden voor werkelijke duur, fatigue na afloop en skip-reden toevoegen.", "UI", "Bijbehorende velden verschijnen contextueel per status."],
            ["ST-403", "Ondersteuning voor ongeplande activiteiten toevoegen.", "Build", "Ongeplande activiteit telt mee in werkelijke totalen."],
            ["ST-404", "Dagoverzicht bouwen met gepland versus uitgevoerd en statusverdeling.", "UI", "Gebruiker ziet de dag samengevat in één scherm."],
            ["ST-405", "Server-side aggregatie voor dagtotalen en eenvoudige samenvatting implementeren.", "Logic", "Dagtotalen blijven consistent met individuele records."],
        ],
    )

    p(doc, "10. Epic 06: Weekoverzicht en inzichten", "Heading 1")
    p(doc, "Doel: terugkijken op patronen zonder de wellness-guardrails te verlaten.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-501", "Weekoverzichtspagina ontwerpen en bouwen.", "UI", "Gebruiker kan per week terugkijken."],
            ["ST-502", "Aggregaties voor gemiddelde energie en budget-adherence bouwen.", "Logic", "Weekstatistieken zijn herleidbaar en testbaar."],
            ["ST-503", "Skip-patronen per activiteit of reden zichtbaar maken.", "Logic/UI", "Patronen worden alleen bij voldoende data getoond."],
            ["ST-504", "Insightregels met minimale datadrempels definiëren.", "Safety/Logic", "Geen patroonclaim zonder expliciete guardrails."],
            ["ST-505", "Tekstuele insightcopy toetsen op niet-medische formulering.", "Content", "Alle teksten blijven binnen de wellness-positionering."],
        ],
    )

    p(doc, "11. Epic 07: Reflectie en reminders", "Heading 1")
    p(doc, "Doel: gebruikers optioneel laten terugblikken na zwaardere dagen.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-601", "Model en flow voor ReflectionCheckIn implementeren.", "Build", "Reflecties kunnen gekoppeld worden aan een eerdere dag."],
            ["ST-602", "Joblogica bouwen die bepaalt welke gebruikers een T+1/T+2 prompt zien.", "Logic/Ops", "Prompts worden niet dubbel of willekeurig aangemaakt."],
            ["ST-603", "Instellingsoptie voor reflectieprompts toevoegen.", "Build", "Gebruiker kan opt-in zelfstandig beheren."],
            ["ST-604", "Korte reflectie-UI bouwen.", "UI", "Prompt voelt licht en niet medisch."],
        ],
    )

    p(doc, "12. Epic 08: Security en operations", "Heading 1")
    p(doc, "Doel: de wellness-first MVP technisch hard genoeg maken voor echte gebruikers.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-701", "Rate limiting op auth- en mutatieroutes toevoegen.", "Security", "Misbruik wordt beperkt op kritieke routes."],
            ["ST-702", "Logging van fouten, loginproblemen en belangrijke mutaties inrichten.", "Ops", "Kerngebeurtenissen zijn herleidbaar."],
            ["ST-703", "Back-up en herstelstrategie voor Supabase documenteren en testen.", "Ops", "Restore-pad is minimaal een keer geoefend of aantoonbaar gevalideerd."],
            ["ST-704", "Secrets- en environmentbeheer voor Vercel en Supabase formaliseren.", "Security/Ops", "Geen secrets in code of onveilige configuratie."],
            ["ST-705", "RLS-policy tests en toegangstests toevoegen.", "QA/Security", "Owner-only model is aantoonbaar afgedwongen."],
        ],
    )

    p(doc, "13. Epic 09: Launch-readiness", "Heading 1")
    p(doc, "Doel: release 1 verantwoord kunnen opleveren.", "Normal")
    table(
        doc,
        ["Story ID", "Omschrijving", "Type", "Definition of done"],
        [
            ["ST-801", "Kernflows handmatig testen op mobiel en desktop.", "QA", "Belangrijkste user journeys zijn geverifieerd."],
            ["ST-802", "Accessibility check uitvoeren op touch targets, contrast en reduced motion.", "QA/UX", "Belangrijkste toegankelijkheidseisen zijn afgevinkt."],
            ["ST-803", "Copy review doen op intended use en non-intended use guardrails.", "Content/Safety", "Geen medische of zorgdossier-taal in release 1."],
            ["ST-804", "DPIA-input en datacatalogus afronden voor de werkelijke MVP-scope.", "Privacy", "Pre-launch privacyartefacten zijn gereed."],
            ["ST-805", "Go-live checklist opstellen met rollback, monitoring en incidentverantwoordelijkheid.", "Ops", "Team weet hoe launch en eerste incidentrespons verloopt."],
        ],
    )

    p(doc, "14. Definition of done op release-niveau", "Heading 1")
    bullets(
        doc,
        [
            "Alle P0-epics zijn functioneel afgerond.",
            "Geen blocking bugs in ochtendcheck-in, planning, evaluatie of dashboardflow.",
            "Owner-only toegang is technisch afgedwongen en getest.",
            "Launchcopy blijft binnen wellness/self-management claims.",
            "Privacy- en securitybasis is gereed voor echte gebruikersintroductie.",
        ],
    )

    p(doc, "15. Bewust níet in deze backlog voor release 1", "Heading 1")
    bullets(
        doc,
        [
            "Viewerrollen, delen met zorgverleners of naasten, en granular sharing.",
            "Habit tracking buiten slaapkwaliteit.",
            "Database-gestuurde vertalingen of extra talen.",
            "AI-inzichten, chatbotfuncties of vrije tekstinterpretatie.",
            "PDF-export, zorgdossierkoppelingen of medical-track features.",
        ],
    )

    p(doc, "16. Aanbevolen vertaling naar projectsturing", "Heading 1")
    numbered(
        doc,
        [
            "Gebruik epics als hoofdstructuur in je projectboard.",
            "Gebruik de story-ID’s als eerste backlogbasis; verfijn ze later in kleinere technische taken.",
            "Koppel iedere story terug aan de bestaande FR-, PRIV-, SEC- en SAFE-eisen waar relevant.",
            "Plan EPIC-08 en EPIC-09 niet als eindklus, maar parallel aan de featurebouw.",
            "Gebruik dit document als leidraad voor scopebewaking: alles wat hier niet in staat, komt niet ongemerkt mee in release 1.",
        ],
    )

    set_footer(doc, f"{PRODUCT_NAME} Implementatieplan en Backlog v0.1")
    doc.save(BASE_DIR / "inspannings-monitor-06-implementatieplan-en-backlog-v01.docx")


def main() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    build_productkader()
    build_functionele_specificatie()
    build_privacy_security_safety()
    build_roadmap()
    build_technische_architectuur()
    build_implementatieplan_backlog()


if __name__ == "__main__":
    main()
